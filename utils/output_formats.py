import os
from docx import Document
import pandas as pd
from graphviz import Digraph

def create_docx(title, content, output_path):
    """Create .docx file"""
    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in content:
        doc.add_paragraph(str(paragraph))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

def create_xlsx(data, output_path):
    """Create .xlsx file"""
    df = pd.DataFrame(data)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    df.to_excel(output_path, index=False)
    return output_path

def create_image(graph, output_path, format="png"):
    """Create image file from Graphviz"""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    graph.render(output_path, format=format, cleanup=True)
    return f"{output_path}.{format}"

def create_md(content, output_path):
    """Create .md file"""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path