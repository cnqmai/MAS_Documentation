# Modified design_tasks.py
import os
import logging
import re
from crewai import Task
from textwrap import dedent
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches
import graphviz

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Placeholder for create_quality_gate_task and create_research_task if not imported
def create_quality_gate_task(agent, phase_name, main_output_key, deliverables_description, output_filename):
    return Task(
        description=dedent(f"""
            Thực hiện kiểm tra chất lượng (Quality Gate) cho giai đoạn {phase_name}.
            Đánh giá tính hoàn chỉnh, rõ ràng, nhất quán và tuân thủ của các deliverables sau:
            {deliverables_description}.
            Tạo một báo cáo tóm tắt kết quả kiểm tra, liệt kê các phát hiện, rủi ro, và đề xuất cải thiện.
            Xác nhận giai đoạn có đủ điều kiện để chuyển tiếp hay không.
        """),
        expected_output=f"Báo cáo chất lượng tiếng Việt '{output_filename}' cho giai đoạn {phase_name}.",
        agent=agent,
        context=[],
        callback=None
    )

def create_research_task(agent, phase_name, research_topic, output_filename):
    return Task(
        description=dedent(f"""
            Nghiên cứu sâu về các phương pháp và tiêu chuẩn thiết kế phần mềm, bao gồm kiến trúc, UX/UI, và bảo mật.
            Tổng hợp các kiến thức này vào một báo cáo để hướng dẫn các agent thiết kế.
            --- Context: Chủ đề nghiên cứu: {research_topic}
        """),
        expected_output=f"Báo cáo nghiên cứu tiếng Việt '{output_filename}' tổng hợp các kiến thức nghiên cứu hữu ích về {research_topic}.",
        agent=agent,
        context=[],
        callback=None
    )

# Helper function to process and save DOCX
def process_and_create_design_doc(task_output: str, doc_title: str, output_path: str, diagram_label: str = ""):
    logging.info(f"--- Bắt đầu xử lý output cho {doc_title} ---")
    doc = Document()
    doc.add_heading(doc_title, level=1)

    dot_blocks = re.findall(r'```dot\s*([\s\S]*?)\s*```', task_output)
    text_content = re.sub(r'```dot\s*[\s\S]*?```', '', task_output).strip()

    # Thêm nội dung văn bản
    if text_content:
        for paragraph in text_content.split('\n\n'):
            doc.add_paragraph(paragraph)

    # Xử lý và thêm biểu đồ nếu có
    if dot_blocks:
        for i, dot_code in enumerate(dot_blocks):
            try:
                graph = graphviz.Source(dot_code)
                diagram_filename = f"{os.path.splitext(os.path.basename(output_path))[0]}_diagram_{i}.png"
                diagram_path = os.path.join(os.path.dirname(output_path), diagram_filename)
                graph.render(diagram_path, format='png', cleanup=True)
                doc.add_heading(f'{diagram_label} Diagram {i+1}', level=2)
                doc.add_picture(diagram_path, width=Inches(6))
                logging.info(f"Đã tạo và thêm {diagram_label} Diagram {i+1} vào tài liệu.")
            except Exception as e:
                logging.error(f"Lỗi khi tạo hoặc thêm biểu đồ DOT vào {doc_title}: {e}")

    try:
        doc.save(output_path)
        logging.info(f"Đã lưu {doc_title} tại: {output_path}")
    except Exception as e:
        logging.error(f"Lỗi khi lưu tài liệu {doc_title}: {e}")

# Main function to create design tasks
def create_design_tasks(design_agent, project_manager_agent, researcher_agent, output_base_dir): # MODIFIED: Removed software_architect_agent
    """
    Tạo các task liên quan đến giai đoạn Thiết kế.
    Args:
        design_agent: Agent chính cho Design.
        project_manager_agent, researcher_agent: Các agent chung.
        output_base_dir: Đường dẫn thư mục base.
    Returns:
        list: Danh sách các Task đã tạo.
    """
    logging.info("--- Tạo các task cho Giai đoạn Thiết kế (Phase 3) ---")

    phase_output_dir = os.path.join(output_base_dir, "3_design")
    os.makedirs(phase_output_dir, exist_ok=True)

    # Fetching necessary inputs from shared_memory
    initial_request = shared_memory.get("system_input", "initial_request") or "Yêu cầu ban đầu chưa có."
    srs_content = shared_memory.get("phase_2_requirements", "srs_document") or "Chưa có SRS."
    brd_content = shared_memory.get("phase_2_requirements", "brd_document") or "Chưa có BRD."
    project_plan_summary = shared_memory.get("phase_1_planning", "project_plan_summary_path") or "Chưa có kế hoạch dự án."
    quality_gate_requirements_report = shared_memory.get("phase_2_requirements", "validation_report_path") or "Chưa có báo cáo Quality Gate giai đoạn Yêu cầu."
    
    design_context = dedent(f"""
        Dựa trên các tài liệu và thông tin sau:
        - Yêu cầu ban đầu của người dùng: {initial_request}
        - Tài liệu Yêu cầu Phần mềm (SRS): {srs_content}
        - Tài liệu Yêu cầu Kinh doanh (BRD): {brd_content}
        - Tóm tắt Kế hoạch Dự án: {project_plan_summary}
        - Báo cáo Quality Gate Giai đoạn Yêu cầu: {quality_gate_requirements_report}
    """)

    # --- 1. Task: Research Design Best Practices (Researcher) ---
    research_design_task = create_research_task(
        researcher_agent,
        "Phase 3: Design",
        dedent(f"""
            Nghiên cứu các phương pháp hay nhất (best practices) và tiêu chuẩn ngành trong thiết kế phần mềm,
            bao gồm kiến trúc hệ thống (microservices, monolithic, client-server), thiết kế cơ sở dữ liệu (SQL/NoSQL),
            thiết kế API (RESTful, GraphQL), thiết kế giao diện người dùng (UX/UI principles, design patterns),
            và các khía cạnh bảo mật trong thiết kế.
            Tổng hợp kiến thức này để hỗ trợ các agent thiết kế.
            --- Context: {design_context}
        """),
        "Design_Research_Summary.md"
    )
    research_design_task.callback = lambda output: (\
        logging.info(f"--- Hoàn thành Research Design Best Practices Task ---"),
        write_output(os.path.join(phase_output_dir, "Design_Research_Summary.md"), str(output)),
        shared_memory.set("phase_3_design", "research_summary_path", os.path.join(phase_output_dir, "Design_Research_Summary.md")),
        logging.info(f"Đã lưu Design_Research_Summary.md và cập nhật shared_memory.")
    )


    # --- 2. Task: System Architecture Design (Design Agent) ---
    architecture_task_description = dedent(f"""
        Thiết kế kiến trúc tổng thể của hệ thống (High-Level Design - HLD), bao gồm các thành phần chính,
        mối quan hệ giữa chúng, các công nghệ dự kiến sử dụng (ngôn ngữ lập trình, framework, CSDL, hosting).
        Tạo sơ đồ kiến trúc (ví dụ: Component Diagram, Deployment Diagram đơn giản) bằng cú pháp DOT và mô tả bằng văn bản.
        Tài liệu cần bao gồm cả sơ đồ và mô tả chi tiết.
        --- Context: {design_context}
        --- Research Summary: {shared_memory.get('phase_3_design', 'research_summary_path') or 'Chưa có tóm tắt nghiên cứu.'}
    """)
    architecture_task = Task(
        description=architecture_task_description,
        expected_output="Tài liệu 'System_Architecture.docx' bao gồm mô tả kiến trúc hệ thống (HLD) và sơ đồ kiến trúc (có thể là Component Diagram hoặc Deployment Diagram) sử dụng cú pháp DOT (sẽ được parse thành ảnh).",
        agent=design_agent,
        context=[research_design_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành System Architecture Design Task ---"),
            process_and_create_design_doc(str(output), "Tài liệu Kiến trúc Hệ thống", os.path.join(phase_output_dir, "System_Architecture.docx"), "System Architecture"),
            shared_memory.set("phase_3_design", "system_architecture_path", os.path.join(phase_output_dir, "System_Architecture.docx")),
            logging.info(f"Đã lưu System_Architecture.docx và cập nhật shared_memory.")
        )
    )

    # --- 3. Task: Functional Design Specification (Design Agent) ---
    design_spec_task_description = dedent(f"""
        Chuyển đổi các yêu cầu chức năng từ SRS/BRD thành các bản đặc tả thiết kế chi tiết (Functional Design Specification - FDS).
        Mô tả cách các tính năng sẽ được hiện thực hóa ở cấp độ thiết kế, bao gồm luồng người dùng, giao diện chính và logic nghiệp vụ.
        --- Context: {design_context}
        --- System Architecture: {shared_memory.get('phase_3_design', 'system_architecture_path') or 'Chưa có kiến trúc hệ thống.'}
    """)
    design_spec_task = Task(
        description=design_spec_task_description,
        expected_output="Tài liệu 'Functional_Design_Specification.docx' mô tả chi tiết cách các chức năng được thiết kế.",
        agent=design_agent,
        context=[architecture_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành Functional Design Specification Task ---"),
            process_and_create_design_doc(str(output), "Đặc tả Thiết kế Chức năng", os.path.join(phase_output_dir, "Functional_Design_Specification.docx")),
            shared_memory.set("phase_3_design", "functional_design_spec_path", os.path.join(phase_output_dir, "Functional_Design_Specification.docx")),
            logging.info(f"Đã lưu Functional_Design_Specification.docx và cập nhật shared_memory.")
        )
    )

    # --- 4. Task: Data Flow Diagram (DFD) Creation (Design Agent) ---
    dfd_task_description = dedent(f"""
        Tạo sơ đồ luồng dữ liệu (Data Flow Diagram - DFD) để hình dung cách dữ liệu di chuyển qua hệ thống.
        Sử dụng cú pháp DOT để vẽ DFD (ít nhất là Context Diagram và Level 0 DFD).
        Cung cấp cả mã DOT và giải thích chi tiết về các thực thể, tiến trình và kho dữ liệu.
        --- Context: {design_context}
        --- Functional Design Specification: {shared_memory.get('phase_3_design', 'functional_design_spec_path') or 'Chưa có FDS.'}
    """)
    dfd_task = Task(
        description=dfd_task_description,
        expected_output="Tài liệu 'Data_Flow_Diagram.md' chứa mã DOT cho DFD (Context, Level 0) và mô tả chi tiết.",
        agent=design_agent,
        context=[design_spec_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành Data Flow Diagram Creation Task ---"),
            write_output(os.path.join(phase_output_dir, "Data_Flow_Diagram.md"), str(output)),
            shared_memory.set("phase_3_design", "data_flow_diagram_md_path", os.path.join(phase_output_dir, "Data_Flow_Diagram.md")),
            logging.info(f"Đã lưu Data_Flow_Diagram.md và cập nhật shared_memory.")
        )
    )

    # --- 5. Task: Sequence Diagram for Key Processes (Design Agent) ---
    sequence_task_description = dedent(f"""
        Chọn ít nhất 2-3 quy trình nghiệp vụ quan trọng hoặc luồng người dùng phức tạp và tạo sơ đồ trình tự (Sequence Diagram) cho chúng.
        Sử dụng cú pháp DOT để vẽ sơ đồ trình tự.
        Cung cấp cả mã DOT và giải thích chi tiết về các đối tượng và tương tác.
        --- Context: {design_context}
        --- Functional Design Specification: {shared_memory.get('phase_3_design', 'functional_design_spec_path') or 'Chưa có FDS.'}
    """)
    sequence_task = Task(
        description=sequence_task_description,
        expected_output="Tài liệu 'Sequence_Diagrams.md' chứa mã DOT cho các sơ đồ trình tự và mô tả chi tiết.",
        agent=design_agent,
        context=[design_spec_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành Sequence Diagram for Key Processes Task ---"),
            write_output(os.path.join(phase_output_dir, "Sequence_Diagrams.md"), str(output)),
            shared_memory.set("phase_3_design", "sequence_diagrams_md_path", os.path.join(phase_output_dir, "Sequence_Diagrams.md")),
            logging.info(f"Đã lưu Sequence_Diagrams.md và cập nhật shared_memory.")
        )
    )

    # --- 6. Task: Database Design (Design Agent) ---
    db_task_description = dedent(f"""
        Thiết kế cấu trúc cơ sở dữ liệu chi tiết, bao gồm các bảng, các trường, kiểu dữ liệu, khóa chính, khóa ngoại, và mối quan hệ.
        Xem xét tính toàn vẹn dữ liệu, hiệu suất và khả năng mở rộng.
        Tạo mô hình cơ sở dữ liệu (ví dụ: ERD) bằng cú pháp DOT và mô tả bảng/trường chi tiết.
        --- Context: {design_context}
        --- Functional Design Specification: {shared_memory.get('phase_3_design', 'functional_design_spec_path') or 'Chưa có FDS.'}
    """)
    db_task = Task(
        description=db_task_description,
        expected_output="Tài liệu 'Database_Design_Document.docx' bao gồm thiết kế CSDL chi tiết và ERD.",
        agent=design_agent,
        context=[design_spec_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành Database Design Task ---"),
            process_and_create_design_doc(str(output), "Tài liệu Thiết kế Cơ sở Dữ liệu", os.path.join(phase_output_dir, "Database_Design_Document.docx"), "ERD"),
            shared_memory.set("phase_3_design", "database_design_doc_path", os.path.join(phase_output_dir, "Database_Design_Document.docx")),
            logging.info(f"Đã lưu Database_Design_Document.docx và cập nhật shared_memory.")
        )
    )

    # --- 7. Task: API Design (Design Agent) ---
    api_task_description = dedent(f"""
        Thiết kế các giao diện lập trình ứng dụng (API) nếu hệ thống có các thành phần cần giao tiếp nội bộ hoặc với bên ngoài.
        Xác định các endpoints, phương thức HTTP, định dạng dữ liệu (request/response, ví dụ JSON schema), và mã trạng thái.
        Sử dụng định dạng YAML hoặc Markdown để mô tả API.
        --- Context: {design_context}
        --- System Architecture: {shared_memory.get('phase_3_design', 'system_architecture_path') or 'Chưa có kiến trúc hệ thống.'}
        --- Functional Design Specification: {shared_memory.get('phase_3_design', 'functional_design_spec_path') or 'Chưa có FDS.'}
    """)
    api_task = Task(
        description=api_task_description,
        expected_output="Tài liệu 'API_Design_Document.md' (hoặc .yaml) mô tả chi tiết các API của hệ thống.",
        agent=design_agent,
        context=[architecture_task, design_spec_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành API Design Task ---"),
            # Giả định output sẽ là YAML hoặc Markdown có thể ghi trực tiếp
            write_output(os.path.join(phase_output_dir, "API_Design_Document.md"), str(output)),
            shared_memory.set("phase_3_design", "api_design_doc_path", os.path.join(phase_output_dir, "API_Design_Document.md")),
            logging.info(f"Đã lưu API_Design_Document.md và cập nhật shared_memory.")
        )
    )

    # --- 8. Task: User Interface/User Experience (UI/UX) Design (Design Agent) ---
    # NOTE: This assumes a simplified UI/UX output. In a real project, this would involve wireframes, mockups, etc.
    ui_ux_task_description = dedent(f"""
        Thiết kế giao diện người dùng (UI) và trải nghiệm người dùng (UX) cho các phần mềm liên quan đến tương tác trực tiếp với người dùng.
        Mô tả các nguyên tắc thiết kế, bố cục màn hình chính, và luồng tương tác người dùng.
        Bạn có thể sử dụng mô tả văn bản hoặc các khối mã giả (pseudo-code/markup) để trình bày ý tưởng UI.
        --- Context: {design_context}
        --- Functional Design Specification: {shared_memory.get('phase_3_design', 'functional_design_spec_path') or 'Chưa có FDS.'}
    """)
    ui_ux_task = Task(
        description=ui_ux_task_description,
        expected_output="Tài liệu 'UI_UX_Design_Principles.md' mô tả các nguyên tắc và bố cục UI/UX chính.",
        agent=design_agent,
        context=[design_spec_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành UI/UX Design Task ---"),
            write_output(os.path.join(phase_output_dir, "UI_UX_Design_Principles.md"), str(output)),
            shared_memory.set("phase_3_design", "ui_ux_design_path", os.path.join(phase_output_dir, "UI_UX_Design_Principles.md")),
            logging.info(f"Đã lưu UI_UX_Design_Principles.md và cập nhật shared_memory.")
        )
    )

    # --- 9. Task: Security Architecture Design (Design Agent) ---
    security_arch_task_description = dedent(f"""
        Thiết kế kiến trúc bảo mật cho hệ thống, xác định các biện pháp bảo mật ở các lớp khác nhau (ứng dụng, dữ liệu, mạng).
        Bao gồm các cơ chế xác thực, ủy quyền, mã hóa, quản lý phiên và xử lý lỗi liên quan đến bảo mật.
        --- Context: {design_context}
        --- System Architecture: {shared_memory.get('phase_3_design', 'system_architecture_path') or 'Chưa có kiến trúc hệ thống.'}
    """)
    security_arch_task = Task(
        description=security_arch_task_description,
        expected_output="Tài liệu 'Security_Architecture_Design.md' mô tả chi tiết kiến trúc bảo mật.",
        agent=design_agent,
        context=[architecture_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành Security Architecture Design Task ---"),
            write_output(os.path.join(phase_output_dir, "Security_Architecture_Design.md"), str(output)),
            shared_memory.set("phase_3_design", "security_architecture_path", os.path.join(phase_output_dir, "Security_Architecture_Design.md")),
            logging.info(f"Đã lưu Security_Architecture_Design.md và cập nhật shared_memory.")
        )
    )

    # --- 10. Task: High-Level Design (HLD) Document (Project Manager / Design Agent) ---
    hld_task_description = dedent(f"""
        Tổng hợp tất cả các quyết định thiết kế cấp cao từ các task trước (kiến trúc, các mô-đun chính, công nghệ)
        vào một tài liệu High-Level Design (HLD) hoàn chỉnh.
        HLD phải cung cấp cái nhìn tổng quan về hệ thống và cách các thành phần tương tác.
        --- Context: {design_context}
        --- System Architecture: {shared_memory.get('phase_3_design', 'system_architecture_path') or 'Chưa có kiến trúc hệ thống.'}
        --- Functional Design Specification: {shared_memory.get('phase_3_design', 'functional_design_spec_path') or 'Chưa có FDS.'}
        --- Database Design: {shared_memory.get('phase_3_design', 'database_design_doc_path') or 'Chưa có thiết kế CSDL.'}
    """)
    hld_task = Task(
        description=hld_task_description,
        expected_output="Tài liệu 'High_Level_Design.docx' tổng hợp các thông tin thiết kế cấp cao.",
        agent=design_agent, # Could be Project Manager too, but Design Agent fits synthesis.
        context=[architecture_task, design_spec_task, db_task, api_task, security_arch_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành High-Level Design Document Task ---"),
            process_and_create_design_doc(str(output), "Tài liệu Thiết kế Cấp cao (HLD)", os.path.join(phase_output_dir, "High_Level_Design.docx")),
            shared_memory.set("phase_3_design", "hld_document_path", os.path.join(phase_output_dir, "High_Level_Design.docx")),
            logging.info(f"Đã lưu High_Level_Design.docx và cập nhật shared_memory.")
        )
    )

    # --- 11. Task: Low-Level Design (LLD) Document (Design Agent) ---
    lld_task_description = dedent(f"""
        Tạo tài liệu Low-Level Design (LLD) cho các mô-đun hoặc tính năng quan trọng.
        LLD đi sâu vào chi tiết về cấu trúc lớp, các lớp, các phương thức, các thuật toán, và cấu trúc dữ liệu cụ thể.
        Chọn một hoặc hai mô-đun quan trọng để thực hiện LLD mẫu.
        --- Context: {design_context}
        --- High-Level Design: {shared_memory.get('phase_3_design', 'hld_document_path') or 'Chưa có HLD.'}
    """)
    lld_task = Task(
        description=lld_task_description,
        expected_output="Tài liệu 'Low_Level_Design.docx' chi tiết thiết kế cấp thấp cho các mô-đun chọn lọc.",
        agent=design_agent,
        context=[hld_task],
        callback=lambda output: (\
            logging.info(f"--- Hoàn thành Low-Level Design Document Task ---"),
            process_and_create_design_doc(str(output), "Tài liệu Thiết kế Cấp thấp (LLD)", os.path.join(phase_output_dir, "Low_Level_Design.docx")),
            shared_memory.set("phase_3_design", "lld_document_path", os.path.join(phase_output_dir, "Low_Level_Design.docx")),
            logging.info(f"Đã lưu Low_Level_Design.docx và cập nhật shared_memory.")
        )
    )

    # --- 12. Task: Design Review Report (Project Manager) ---
    # (Quality Gate for Design Phase - uses quality_gate_tasks.py)
    quality_gate_design_callback = lambda output: (\
        logging.info(f"--- Hoàn thành Quality Gate Design Task ---"),
        write_output(os.path.join(phase_output_dir, "Validation_Report_Phase_3.md"), str(output)),
        shared_memory.set("phase_3_design", "validation_report_path", os.path.join(phase_output_dir, "Validation_Report_Phase_3.md")),
        logging.info(f"Đã lưu Validation_Report_Phase_3.md và cập nhật shared_memory.")
    )

    quality_gate_design_task = create_quality_gate_task(
        project_manager_agent,
        "Phase 3: Design",
        "lld", # Key for content to validate (could be hld or lld)
        dedent("""
        System_Architecture.docx, Functional_Design_Specification.docx, Data_Flow_Diagram.md,
        Sequence_Diagrams.md, Database_Design_Document.docx, API_Design_Document.md,
        UI_UX_Design_Principles.md, Security_Architecture_Design.md, High_Level_Design.docx, Low_Level_Design.docx
        """),
        "Validation_Report_Phase_3.md"
    )
    quality_gate_design_task.context = [
        design_spec_task,
        architecture_task,
        dfd_task,
        sequence_task,
        db_task,
        api_task,
        security_arch_task,
        hld_task,
        lld_task,
        ui_ux_task,
        research_design_task
    ]
    quality_gate_design_task.callback = quality_gate_design_callback

    return [
        research_design_task,
        architecture_task,
        design_spec_task,
        dfd_task,
        sequence_task,
        db_task,
        api_task,
        ui_ux_task, # Added UI/UX task
        security_arch_task,
        hld_task,
        lld_task,
        quality_gate_design_task
    ]