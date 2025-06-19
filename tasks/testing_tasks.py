import os
import logging
import re
from crewai import Task
from textwrap import dedent
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches
import graphviz

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def read_file_content(file_path):
    """Đọc nội dung file nếu tồn tại."""
    if file_path and file_path != "N/A" and os.path.exists(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logging.warning(f"Không thể đọc file {file_path}: {e}")
            return f"N/A (Lỗi đọc file: {e})"
    return "N/A (File không tồn tại)"

def create_testing_tasks(testing_agent, project_manager_agent, output_base_dir):
    """
    Tạo các task liên quan đến giai đoạn Kiểm thử.
    Args:
        testing_agent: Agent chính cho Testing.
        project_manager_agent: Agent cho Audit và Quality Gate.
        output_base_dir: Đường dẫn thư mục base.
    Returns:
        list: Danh sách các Task đã tạo.
    """
    # Tạo thư mục con cho Phase 5 Testing
    phase_output_dir = os.path.join(output_base_dir, "5_testing")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đảm bảo thư mục output cho Phase 5: Testing tồn tại tại: {phase_output_dir}")

    # Lấy các input từ shared_memory với fallback
    srs_content = shared_memory.get("phase_2_requirements", "srs_document") or "N/A"
    use_case_doc_path = shared_memory.get("phase_2_requirements", "use_case_document_path") or "N/A"
    use_case_diagram_path = shared_memory.get("phase_2_requirements", "use_case_diagram_path") or "N/A"
    nfrd_content = shared_memory.get("phase_2_requirements", "nfrd_document") or "N/A"
    project_plan_xml_path = shared_memory.get("phase_1_planning", "project_plan_xml_path") or "N/A"
    hld_doc_path = shared_memory.get("phase_3_design", "hld_document_path") or "N/A"
    source_code_path = shared_memory.get("phase_4_development", "source_code_doc_template") or "N/A"
    coding_standards_path = shared_memory.get("phase_4_development", "coding_standards_document_path") or "N/A"

    # Context string cho các task
    testing_context_str = dedent(f"""
        Dựa trên các tài liệu đã có:
        - Yêu cầu phần mềm (SRS): {srs_content[:500]}... (hoặc từ file: {use_case_doc_path})
        - Sơ đồ Use Case: {use_case_diagram_path}
        - Yêu cầu phi chức năng (NFRD): {nfrd_content[:500]}...
        - Kế hoạch dự án: {read_file_content(project_plan_xml_path)}
        - Tài liệu thiết kế (HLD): {hld_doc_path}
        - Mã nguồn: {source_code_path}
        - Tiêu chuẩn mã hóa: {coding_standards_path}
    """)

    # --- 1. Task: QA Checklist Creation (testing_agent) ---
    def qa_checklist_callback(output):
        logging.info("--- Hoàn thành QA Checklist Creation Task ---")
        file_path = os.path.join(phase_output_dir, "QA_Checklist.md")
        write_output(file_path, str(output))
        shared_memory.set("phase_5_testing", "qa_checklist_path", file_path)
        logging.info("Đã lưu QA_Checklist.md và cập nhật shared_memory.")

    qa_checklist_task = Task(
        description=dedent(f"""
            Tạo một danh sách kiểm tra chất lượng (QA Checklist) chi tiết dựa trên Tài liệu Yêu cầu Phần mềm (SRS),
            Yêu cầu Phi Chức năng (NFRD), tài liệu thiết kế, và các tiêu chuẩn ngành liên quan đến chất lượng phần mềm.
            Danh sách này sẽ được sử dụng để hướng dẫn toàn bộ quá trình kiểm thử.
            Tập trung vào các khía cạnh: tính năng, hiệu năng, bảo mật, khả năng sử dụng, và khả năng tương thích.
            Đầu ra phải là một tài liệu Markdown rõ ràng, có cấu trúc tốt.
            {testing_context_str}
        """),
        expected_output="Một tài liệu Markdown tiếng Việt có cấu trúc tốt 'QA_Checklist.md' bao gồm các hạng mục kiểm tra chi tiết.",
        agent=testing_agent,
        callback=qa_checklist_callback
    )

    # --- 2. Task: Test Plan Document Creation (testing_agent) ---
    def test_plan_callback(output):
        logging.info("--- Hoàn thành Test Plan Document Creation Task ---")
        file_path = os.path.join(phase_output_dir, "Test_Plan.docx")
        # Extract text and DOT code
        dot_blocks = re.findall(r'```dot\s*([\s\S]*?)\s*```', str(output))
        text_content = re.sub(r'```dot\s*[\s\S]*?```', '', str(output)).strip()

        doc = Document()
        doc.add_heading('Kế hoạch Kiểm thử (Test Plan)', level=1)
        doc.add_paragraph(text_content)

        if dot_blocks:
            dot_code = dot_blocks[0]
            graph_path = os.path.join(phase_output_dir, "Test_Strategy_Diagram")
            try:
                src = graphviz.Source(dot_code, format='png')
                src.render(graph_path, view=False, cleanup=True)
                doc.add_heading('Chiến lược Kiểm thử', level=2)
                doc.add_picture(graph_path + '.png', width=Inches(6))
                logging.info("Đã tạo và nhúng sơ đồ chiến lược kiểm thử.")
            except Exception as e:
                logging.error(f"Lỗi khi tạo sơ đồ DOT cho Test Plan: {e}")
                doc.add_paragraph(f"Không thể tạo sơ đồ chiến lược kiểm thử. Mã DOT: {dot_code}")

        doc.save(file_path)
        shared_memory.set("phase_5_testing", "test_plan_path", file_path)
        logging.info("Đã lưu Test_Plan.docx và cập nhật shared_memory.")

    test_plan_task = Task(
        description=dedent(f"""
            Xây dựng một Tài liệu Kế hoạch Kiểm thử (Test Plan) toàn diện, bao gồm:
            - Chiến lược kiểm thử: loại hình kiểm thử (functional, non-functional, security, performance), tiếp cận.
            - Phạm vi kiểm thử.
            - Môi trường kiểm thử.
            - Tiêu chí chấp nhận và kết thúc kiểm thử.
            - Lịch trình và tài nguyên.
            - Vẽ một sơ đồ DOT mô tả luồng/chiến lược kiểm thử chính.
            Sử dụng thông tin từ SRS, NFRD, tài liệu thiết kế và QA Checklist.
            Đảm bảo output được định dạng tốt để ghi vào file .docx.
            {testing_context_str}
            --- QA Checklist: {read_file_content(shared_memory.get("phase_5_testing", "qa_checklist_path") or "N/A")}
        """),
        expected_output="Một tài liệu Word tiếng Việt 'Test_Plan.docx' chứa kế hoạch kiểm thử chi tiết và một sơ đồ DOT cho chiến lược kiểm thử.",
        agent=testing_agent,
        context=[qa_checklist_task],
        callback=test_plan_callback
    )

    # --- 3. Task: Test Case Development (testing_agent) ---
    def test_case_callback(output):
        logging.info("--- Hoàn thành Test Case Development Task ---")
        file_path = os.path.join(phase_output_dir, "Test_Cases.xlsx")
        write_output(file_path, str(output))  # Giả định output là bảng có thể parse thành XLSX
        shared_memory.set("phase_5_testing", "test_cases_path", file_path)
        logging.info("Đã lưu Test_Cases.xlsx và cập nhật shared_memory.")

    test_case_task = Task(
        description=dedent(f"""
            Phát triển các trường hợp kiểm thử (Test Cases) chi tiết dựa trên Tài liệu Yêu cầu Phần mềm (SRS),
            Yêu cầu Phi Chức năng (NFRD), Tài liệu Thiết kế (HLD), và Kế hoạch Kiểm thử.
            Mỗi Test Case cần bao gồm: ID, Mô tả, Điều kiện tiên quyết, Các bước thực hiện, Kết quả mong đợi, Mức độ ưu tiên.
            Tạo các test case cho cả kiểm thử chức năng và phi chức năng.
            Cung cấp output dưới dạng bảng, có thể chuyển đổi thành file Excel.
            {testing_context_str}
            --- Test Plan: {read_file_content(shared_memory.get("phase_5_testing", "test_plan_path") or "N/A")}
        """),
        expected_output="Một tài liệu tiếng Việt 'Test_Cases.xlsx' chứa danh sách các trường hợp kiểm thử chi tiết ở định dạng bảng.",
        agent=testing_agent,
        context=[qa_checklist_task, test_plan_task],
        callback=test_case_callback
    )

    # --- 4. Task: Test Execution and Defect Reporting (testing_agent) ---
    def test_exec_callback(output):
        logging.info("--- Hoàn thành Test Execution and Defect Reporting Task ---")
        file_path = os.path.join(phase_output_dir, "Test_Execution_Report.md")
        write_output(file_path, str(output))
        shared_memory.set("phase_5_testing", "test_execution_report_path", file_path)
        logging.info("Đã lưu Test_Execution_Report.md và cập nhật shared_memory.")

    test_exec_task = Task(
        description=dedent(f"""
            Giả lập việc thực thi các trường hợp kiểm thử đã tạo và báo cáo kết quả.
            Xác định các lỗi (defects) nếu có và ghi lại chúng dưới dạng Báo cáo Lỗi (Defect Report)
            với các thông tin: ID lỗi, Mô tả, Các bước tái hiện, Mức độ ưu tiên, Mức độ nghiêm trọng, Trạng thái.
            Đưa ra một bản tóm tắt về tiến độ kiểm thử và tỷ lệ lỗi.
            Giả định bạn có thể truy cập và chạy mã nguồn để kiểm thử (trong môi trường ảo).
            {testing_context_str}
            --- Test Cases: {read_file_content(shared_memory.get("phase_5_testing", "test_cases_path") or "N/A")}
            --- Source Code: {read_file_content(source_code_path)}
        """),
        expected_output="Một tài liệu Markdown tiếng Việt 'Test_Execution_Report.md' tóm tắt kết quả thực thi kiểm thử và danh sách các lỗi (nếu có).",
        agent=testing_agent,
        context=[qa_checklist_task, test_plan_task, test_case_task],
        callback=test_exec_callback
    )

    # --- 5. Task: Security and Performance Testing (testing_agent) ---
    def security_perf_test_callback(output):
        logging.info("--- Hoàn thành Security and Performance Testing Task ---")
        file_path = os.path.join(phase_output_dir, "Security_Performance_Test_Report.md")
        write_output(file_path, str(output))
        shared_memory.set("phase_5_testing", "security_performance_report_path", file_path)
        logging.info("Đã lưu Security_Performance_Test_Report.md và cập nhật shared_memory.")

    security_perf_test_task = Task(
        description=dedent(f"""
            Thực hiện đánh giá cơ bản về bảo mật (ví dụ: các lỗ hổng OWASP Top 10 phổ biến)
            và kiểm thử hiệu năng (ví dụ: tải, stress test) cho hệ thống.
            Sử dụng NFRD để xác định các yêu cầu bảo mật và hiệu năng.
            Giả lập các công cụ kiểm thử tự động và báo cáo các phát hiện quan trọng.
            {testing_context_str}
            --- Test Plan: {read_file_content(shared_memory.get("phase_5_testing", "test_plan_path") or "N/A")}
            --- Test Cases: {read_file_content(shared_memory.get("phase_5_testing", "test_cases_path") or "N/A")}
        """),
        expected_output="Một tài liệu Markdown tiếng Việt 'Security_Performance_Test_Report.md' tóm tắt các phát hiện về bảo mật và hiệu năng.",
        agent=testing_agent,
        context=[qa_checklist_task, test_plan_task, test_case_task],
        callback=security_perf_test_callback
    )

    # --- 6. Task: Quality Assurance Audit (project_manager_agent) ---
    def audit_callback(output):
        logging.info("--- Hoàn thành Quality Assurance Audit Task ---")
        file_path = os.path.join(phase_output_dir, "Quality_Assurance_Audit_Report.md")
        write_output(file_path, str(output))
        shared_memory.set("phase_5_testing", "qa_audit_report_path", file_path)
        logging.info("Đã lưu Quality_Assurance_Audit_Report.md và cập nhật shared_memory.")

    audit_task = Task(
        description=dedent(f"""
            Thực hiện một cuộc kiểm toán chất lượng (QA Audit) độc lập.
            Đánh giá toàn bộ quy trình kiểm thử, bao gồm kế hoạch kiểm thử, trường hợp kiểm thử,
            báo cáo thực thi và báo cáo lỗi, để đảm bảo chúng tuân thủ các tiêu chuẩn dự án và ngành.
            Xác định các điểm mạnh và các khu vực cần cải thiện trong quy trình QA.
            {testing_context_str}
            --- QA Checklist: {read_file_content(shared_memory.get("phase_5_testing", "qa_checklist_path") or "N/A")}
            --- Test Plan: {read_file_content(shared_memory.get("phase_5_testing", "test_plan_path") or "N/A")}
            --- Test Cases: {read_file_content(shared_memory.get("phase_5_testing", "test_cases_path") or "N/A")}
            --- Test Execution Report: {read_file_content(shared_memory.get("phase_5_testing", "test_execution_report_path") or "N/A")}
            --- Security/Performance Report: {read_file_content(shared_memory.get("phase_5_testing", "security_performance_report_path") or "N/A")}
        """),
        expected_output="Một tài liệu Markdown tiếng Việt 'Quality_Assurance_Audit_Report.md' báo cáo kết quả kiểm toán QA.",
        agent=project_manager_agent,
        context=[qa_checklist_task, test_plan_task, test_case_task, test_exec_task, security_perf_test_task],
        callback=audit_callback
    )

    # --- 7. Task: Test Management Reporting (project_manager_agent) ---
    def test_management_callback(output):
        logging.info("--- Hoàn thành Test Management Reporting Task ---")
        file_path = os.path.join(phase_output_dir, "Test_Management_Report.docx")
        doc = Document()
        doc.add_heading('Báo cáo Quản lý Kiểm thử', level=1)
        doc.add_paragraph(str(output))
        doc.save(file_path)
        shared_memory.set("phase_5_testing", "test_management_report_path", file_path)
        logging.info("Đã lưu Test_Management_Report.docx và cập nhật shared_memory.")

    test_management_task = Task(
        description=dedent(f"""
            Tổng hợp tất cả các báo cáo kiểm thử (Test Execution Report, Security/Performance Report, QA Audit Report)
            thành một báo cáo quản lý tổng thể cho Project Manager.
            Báo cáo này cần cung cấp cái nhìn tổng quan về tình trạng chất lượng sản phẩm, các rủi ro còn tồn đọng,
            và khuyến nghị để đưa ra quyết định "Go/No-Go" cho giai đoạn tiếp theo.
            {testing_context_str}
            --- Test Execution Report: {read_file_content(shared_memory.get("phase_5_testing", "test_execution_report_path") or "N/A")}
            --- Security/Performance Report: {read_file_content(shared_memory.get("phase_5_testing", "security_performance_report_path") or "N/A")}
            --- QA Audit Report: {read_file_content(shared_memory.get("phase_5_testing", "qa_audit_report_path") or "N/A")}
        """),
        expected_output="Một tài liệu Word tiếng Việt 'Test_Management_Report.docx' tóm tắt toàn bộ tình hình kiểm thử và đề xuất quyết định.",
        agent=project_manager_agent,
        context=[test_exec_task, security_perf_test_task, audit_task],
        callback=test_management_callback
    )

    # --- 8. Task: Quality Gate Review (project_manager_agent) ---
    def quality_gate_callback(output):
        logging.info("--- Hoàn thành Quality Gate Review Task for Testing ---")
        file_path = os.path.join(phase_output_dir, "Validation_Report_Phase_5.md")
        write_output(file_path, str(output))
        shared_memory.set("phase_5_testing", "validation_report_path", file_path)
        logging.info("Đã lưu Validation_Report_Phase_5.md và cập nhật shared_memory.")

    quality_gate_testing_task = Task(
        description=dedent(f"""
            Thực hiện Quality Gate cuối cùng cho giai đoạn Kiểm thử.
            Đánh giá kỹ lưỡng tất cả tài liệu và báo cáo kiểm thử (QA Checklist, Test Plan, Test Cases,
            Test Execution Report, Security/Performance Report, QA Audit Report, Test Management Report).
            Kiểm tra tính đầy đủ, chính xác, nhất quán và tuân thủ các tiêu chuẩn chất lượng.
            Xác nhận giai đoạn Kiểm thử có đủ điều kiện để chuyển tiếp sang giai đoạn Triển khai và Giao nộp hay không.
            {testing_context_str}
            --- QA Checklist: {read_file_content(shared_memory.get("phase_5_testing", "qa_checklist_path") or "N/A")}
            --- Test Plan: {read_file_content(shared_memory.get("phase_5_testing", "test_plan_path") or "N/A")}
            --- Test Cases: {read_file_content(shared_memory.get("phase_5_testing", "test_cases_path") or "N/A")}
            --- Test Execution Report: {read_file_content(shared_memory.get("phase_5_testing", "test_execution_report_path") or "N/A")}
            --- Security/Performance Report: {read_file_content(shared_memory.get("phase_5_testing", "security_performance_report_path") or "N/A")}
            --- QA Audit Report: {read_file_content(shared_memory.get("phase_5_testing", "qa_audit_report_path") or "N/A")}
            --- Test Management Report: {read_file_content(shared_memory.get("phase_5_testing", "test_management_report_path") or "N/A")}
        """),
        expected_output="Một tài liệu Markdown tiếng Việt 'Validation_Report_Phase_5.md' nêu rõ kết quả kiểm tra Quality Gate, lý do và các đề xuất cải thiện.",
        agent=project_manager_agent,
        context=[qa_checklist_task, test_plan_task, test_case_task, test_exec_task,
                 security_perf_test_task, audit_task, test_management_task],
        callback=quality_gate_callback
    )

    return [
        qa_checklist_task,
        test_plan_task,
        test_case_task,
        test_exec_task,
        security_perf_test_task,
        audit_task,
        test_management_task,
        quality_gate_testing_task
    ]