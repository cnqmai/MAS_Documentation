import os
import logging
import re
from crewai import Task
from textwrap import dedent # Import dedent for cleaner multi-line strings
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches
import graphviz

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Helper functions for processing task outputs ---
# Các hàm này sẽ xử lý đầu ra của Task và ghi vào file, cập nhật shared_memory

def process_and_create_requirements_doc(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho Tài liệu Yêu cầu ---")
    doc = Document()
    doc.add_heading('Tài liệu Yêu cầu Phần mềm (SRS)', level=1)

    # Tách các khối mã DOT (giả định Use Case Diagram)
    dot_blocks = re.findall(r'```dot\s*([\s\S]*?)\s*```', task_output)
    text_content = re.sub(r'```dot\s*[\s\S]*?```', '', task_output).strip()

    use_case_text = text_content # Giả định phần còn lại là văn bản Use Cases
    use_case_dot_code = ""

    if len(dot_blocks) >= 1:
        use_case_dot_code = dot_blocks[0]
        logging.info("Đã trích xuất mã DOT cho Use Case Diagram.")

    # Thêm nội dung văn bản Use Case
    if use_case_text:
        doc.add_heading('Các Trường hợp Sử dụng (Use Cases)', level=2)
        doc.add_paragraph(use_case_text)
    else:
        doc.add_paragraph("Không có nội dung Use Case dạng văn bản được tạo.")

    # Đảm bảo thư mục con tồn tại cho output của phase này
    phase_output_dir = os.path.join(output_base_dir_param, "2_requirements")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 2: {phase_output_dir}")

    # Render Use Case Diagram
    if use_case_dot_code:
        try:
            graph_use_case = graphviz.Source(use_case_dot_code, format='png', engine='dot')
            use_case_img_path = os.path.join(phase_output_dir, "Use_Case_Diagram.png")
            graph_use_case.render(use_case_img_path.rsplit('.', 1)[0], view=False, cleanup=True)
            doc.add_heading('Use Case Diagram', level=2)
            doc.add_picture(use_case_img_path, width=Inches(6.0))
            logging.info(f"Đã tạo và chèn Use Case Diagram vào tài liệu: {use_case_img_path}")
            shared_memory.set("phase_2_requirements", "use_case_diagram_path", use_case_img_path)
        except Exception as e:
            logging.error(f"Lỗi khi tạo Use Case Diagram: {e}", exc_info=True)
            doc.add_paragraph(f"Không thể tạo Use Case Diagram do lỗi: {e}\nMã DOT thất bại:\n```dot\n{use_case_dot_code}\n```")
    else:
        doc.add_paragraph("Agent không tạo ra mã DOT cho Use Case Diagram.")

    final_doc_path = os.path.join(phase_output_dir, "Use_Case_Document_with_Diagram.docx")
    doc.save(final_doc_path)
    logging.info(f"Đã lưu tài liệu Use_Case_Document_with_Diagram.docx vào {final_doc_path}")

    # The use_case_task also outputs User_Stories.md. We should separate this.
    # For now, let's assume the text_content above is the user stories or part of it.
    user_stories_path = os.path.join(phase_output_dir, "User_Stories.md")
    write_output(user_stories_path, use_case_text if use_case_text else "No user stories generated.")
    shared_memory.set("phase_2_requirements", "user_stories_path", user_stories_path)
    logging.info(f"Đã lưu User_Stories.md vào {user_stories_path}")

    shared_memory.set("phase_2_requirements", "use_case_document_path", final_doc_path)
    logging.info(f"--- Hoàn thành xử lý output cho Tài liệu Yêu cầu ---")


def create_requirement_tasks(requirement_analyst_agent, project_manager_agent, researcher_agent, output_base_dir):
    """
    Tạo các task liên quan đến giai đoạn Yêu cầu.
    Args:
        requirement_analyst_agent: Agent chính cho Requirements.
        project_manager_agent, researcher_agent: Các agent chung.
        output_base_dir: Đường dẫn thư mục base.
    Returns:
        list: Danh sách các Task đã tạo.
    """
    # Assuming WBS.xlsx and Project_Plan.xml or their content are available via shared_memory or direct input
    # For this example, we'll simulate inputs or assume they're handled upstream.
    # In a real scenario, these would be loaded from files or previous task outputs.
    wbs_content = shared_memory.get("phase_1_planning", "wbs_content") or "Nội dung WBS giả định."
    project_plan_content = shared_memory.get("phase_1_planning", "project_plan") or "Nội dung Project Plan giả định."
    conops_content = shared_memory.get("phase_1_planning", "conops_document") or "Nội dung CONOPS giả định."

    if not project_plan_content:
        logging.warning("Project Plan content missing for requirements tasks.")
        project_plan_content = "Không có tài liệu Project Plan."

    # Tạo thư mục con cho Phase 2 Requirements (nếu chưa có)
    phase_output_dir = os.path.join(output_base_dir, "2_requirements")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 2: {phase_output_dir}")

    # Task: Project Scope Definition (scope_tasks.py)
    # Output: Scope_Requirements_Checklist.xlsx (but generating as markdown for simplicity)
    scope_task = Task(
        description=dedent(f"""
            Định nghĩa rõ ràng phạm vi của dự án, bao gồm các mục tiêu chính, các kết quả mong đợi,
            và các giới hạn của dự án. Tạo ra một checklist các yêu cầu về phạm vi.
            --- Context: Project Plan: {project_plan_content}, WBS: {wbs_content}
        """),
        expected_output="Tài liệu tiếng Việt 'Scope_Requirements_Checklist.md' chi tiết và rõ ràng, định nghĩa phạm vi.",
        agent=requirement_analyst_agent, # Changed to requirement_agent as per table
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Project Scope Definition Task ---"),
            write_output(os.path.join(phase_output_dir, "Scope_Requirements_Checklist.md"), str(output)),
            shared_memory.set("phase_2_requirements", "scope_requirements_checklist", str(output)),
            logging.info(f"Đã lưu Scope_Requirements_Checklist.md và cập nhật shared_memory.")
        )
    )

    # Task: Business Requirements Document (BRD) (brd_tasks.py)
    # Inputs: Scope_Requirements_Checklist.xlsx
    # Output: Business_Requirements_Document.docx
    brd_task = Task(
        description=dedent(f"""
            Dựa trên Project Plan, Project Scope, và checklist yêu cầu phạm vi,
            phát triển Tài liệu Yêu cầu Nghiệp vụ (BRD).
            Tập trung vào các yêu cầu cấp cao từ góc độ kinh doanh và stakeholder, không đi sâu vào chi tiết kỹ thuật.
            --- Project Plan: {project_plan_content}
            --- Project Scope/Scope Requirements Checklist: {scope_task.output.raw_output if scope_task.output else 'Chưa có Scope Requirements Checklist.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Business_Requirements_Document.docx' đầy đủ và có cấu trúc.",
        agent=requirement_analyst_agent,
        context=[scope_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành BRD Task ---"),
            write_output(os.path.join(phase_output_dir, "Business_Requirements_Document.docx"), str(output)),
            shared_memory.set("phase_2_requirements", "brd_document", str(output)),
            logging.info(f"Đã lưu Business_Requirements_Document.docx và cập nhật shared_memory.")
        )
    )

    # Task: BRD Presentation (presentation_tasks.py)
    # Inputs: Business_Requirements_Document.docx
    # Output: BRD_Presentation.pptx (generating as markdown summary)
    brd_presentation_task = Task(
        description=dedent(f"""
            Chuẩn bị một bản trình bày (dưới dạng văn bản markdown, tóm tắt các slide chính)
            về Tài liệu Yêu cầu Nghiệp vụ (BRD) cho các stakeholder.
            Tóm tắt các điểm chính, phạm vi, mục tiêu và các yêu cầu nghiệp vụ cấp cao.
            --- BRD: {brd_task.output.raw_output if brd_task.output else 'Chưa có BRD.'}
        """),
        expected_output="Tài liệu tiếng Việt 'BRD_Presentation.md' tóm tắt nội dung trình bày BRD.",
        agent=requirement_analyst_agent, # Changed to requirement_agent as per table
        context=[brd_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành BRD Presentation Task ---"),
            write_output(os.path.join(phase_output_dir, "BRD_Presentation.md"), str(output)),
            shared_memory.set("phase_2_requirements", "brd_presentation_document", str(output)),
            logging.info(f"Đã lưu BRD_Presentation.md và cập nhật shared_memory.")
        )
    )

    # Task: Software Requirements Specification (SRS) (srs_tasks.py)
    # Inputs: Business_Requirements_Document.docx
    # Output: System_Requirements_Specification.docx
    srs_task = Task(
        description=dedent(f"""
            Từ BRD (có trong context), chi tiết hóa các yêu cầu thành Tài liệu Đặc tả Yêu cầu Phần mềm (SRS).
            Bao gồm các yêu cầu chức năng, phi chức năng, giao diện, và các ràng buộc.
            --- BRD: {brd_task.output.raw_output if brd_task.output else 'Chưa có BRD.'}
        """),
        expected_output="Tài liệu tiếng Việt 'System_Requirements_Specification.docx' chi tiết, bao gồm tất cả các loại yêu cầu.",
        agent=requirement_analyst_agent,
        context=[brd_task], # SRS depends on BRD
        callback=lambda output: (
            logging.info(f"--- Hoàn thành SRS Task ---"),
            write_output(os.path.join(phase_output_dir, "System_Requirements_Specification.docx"), str(output)),
            shared_memory.set("phase_2_requirements", "srs_document", str(output)),
            logging.info(f"Đã lưu System_Requirements_Specification.docx và cập nhật shared_memory.")
        )
    )

    # Task: Functional Requirements Document (FRD) (frd_tasks.py)
    # Inputs: Scope_Requirements_Checklist.xlsx, Business_Requirements_Document.docx
    # Output: Functional_Requirements_Document.docx (F.R.D)
    frd_task = Task(
        description=dedent(f"""
            Dựa trên Scope_Requirements_Checklist và Business_Requirements_Document (có trong context),
            phát triển Tài liệu Yêu cầu Chức năng (FRD).
            Tập trung mô tả chi tiết từng chức năng của hệ thống, bao gồm các kịch bản sử dụng.
            --- Scope Requirements Checklist: {scope_task.output.raw_output if scope_task.output else 'Chưa có Scope Requirements Checklist.'}
            --- BRD: {brd_task.output.raw_output if brd_task.output else 'Chưa có BRD.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Functional_Requirements_Document.docx' mô tả chi tiết các yêu cầu chức năng.",
        agent=requirement_analyst_agent,
        context=[scope_task, brd_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành FRD Task ---"),
            write_output(os.path.join(phase_output_dir, "Functional_Requirements_Document.docx"), str(output)),
            shared_memory.set("phase_2_requirements", "frd_document", str(output)),
            logging.info(f"Đã lưu Functional_Requirements_Document.docx và cập nhật shared_memory.")
        )
    )

    # Task: Use Case Document and Diagram (usecase_tasks.py)
    # Inputs: SRS.docx, CONOPS.docx
    # Outputs: Use_Case_Diagram.png, User_Stories.md (handled by callback)
    use_case_task = Task(
        description=dedent(f"""
            Dựa trên SRS và CONOPS (có trong context), phát triển các Use Case chi tiết và sơ đồ Use Case.
            Mỗi Use Case cần mô tả: tên, mục tiêu, các tác nhân, luồng sự kiện chính và luồng thay thế.
            Bạn phải **tạo mã nguồn Graphviz DOT** để biểu diễn Use Case dưới dạng sơ đồ Use Case.
            Cấu trúc đầu ra của bạn phải là văn bản mô tả Use Cases (sẽ được dùng làm User Stories), sau đó là mã DOT cho Use Case Diagram
            (trong '```dot\\n...\\n```').
            --- SRS: {srs_task.output.raw_output if srs_task.output else 'Chưa có SRS.'}
            --- CONOPS: {conops_content}
        """),
        expected_output=(
            "Một chuỗi văn bản (string) bao gồm:\n"
            "1. Phần mô tả các Use Cases (User Stories).\n"
            "2. Tiếp theo là mã Graphviz DOT cho Use Case Diagram được bọc trong '```dot\\n...\\n```'.\n"
            "Đảm bảo mã DOT đúng cú pháp để có thể render thành hình ảnh."
        ),
        agent=requirement_analyst_agent,
        context=[srs_task], # CONOPS content is passed directly
        callback=lambda output: process_and_create_requirements_doc(str(output), output_base_dir)
    )

    # Task: Requirements Traceability Matrix (RTM) (rtm_tasks.py)
    # Inputs: SRS.docx
    # Output: Requirements_Traceability_Matrix.xlsx (generating as markdown)
    rtm_task = Task(
        description=dedent(f"""
            Xây dựng Ma trận Truy vết Yêu cầu (RTM) để liên kết các yêu cầu từ BRD đến SRS, FRD, NFRD, và Use Cases.
            Đảm bảo mỗi yêu cầu được truy vết và không bị bỏ sót.
            --- SRS: {srs_task.output.raw_output if srs_task.output else 'Chưa có SRS.'}
            --- BRD: {brd_task.output.raw_output if brd_task.output else 'Chưa có BRD.'}
            --- FRD: {frd_task.output.raw_output if frd_task.output else 'Chưa có FRD.'}
            --- NFRD: {shared_memory.get("phase_2_requirements", "nfrd_document") or 'Chưa có NFRD.'}
            --- Use Cases: {use_case_task.output.raw_output if use_case_task.output else 'Chưa có Use Cases.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Requirements_Traceability_Matrix.md' thể hiện mối quan hệ truy vết.",
        agent=requirement_analyst_agent,
        context=[srs_task, brd_task, frd_task, use_case_task], # NFRD content will be needed
        callback=lambda output: (
            logging.info(f"--- Hoàn thành RTM Task ---"),
            write_output(os.path.join(phase_output_dir, "Requirements_Traceability_Matrix.md"), str(output)),
            shared_memory.set("phase_2_requirements", "rtm_document", str(output)),
            logging.info(f"Đã lưu Requirements_Traceability_Matrix.md và cập nhật shared_memory.")
        )
    )

    # NEW TASK: Impact Analysis Report (impact_tasks.py)
    # Inputs: Requirements_Traceability_Matrix.xlsx
    # Output: Change_Impact_Analysis_Report.docx (generating as markdown)
    impact_analysis_task = Task(
        description=dedent(f"""
            Thực hiện phân tích tác động (impact analysis) dựa trên Ma trận Truy vết Yêu cầu (RTM).
            Xác định ảnh hưởng của bất kỳ thay đổi nào đối với các yêu cầu đã được xác định.
            --- RTM: {rtm_task.output.raw_output if rtm_task.output else 'Chưa có RTM.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Change_Impact_Analysis_Report.md' chi tiết về phân tích tác động.",
        agent=requirement_analyst_agent, # Assign to requirement_agent as per table
        context=[rtm_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Change Impact Analysis Task ---"),
            write_output(os.path.join(phase_output_dir, "Change_Impact_Analysis_Report.md"), str(output)),
            shared_memory.set("phase_2_requirements", "change_impact_analysis_report", str(output)),
            logging.info(f"Đã lưu Change_Impact_Analysis_Report.md và cập nhật shared_memory.")
        )
    )

    # Task: Service Level Agreement (SLA) Definition (sla_tasks.py)
    # Inputs: SRS.docx
    # Output: Service_Level_Agreement_Template.docx (generating as markdown)
    sla_task = Task(
        description=dedent(f"""
            Định nghĩa các chỉ số hiệu suất, độ tin cậy, bảo mật và các cam kết SLA (Service Level Agreement) khác
            dựa trên các yêu cầu từ SRS.
            --- SRS: {srs_task.output.raw_output if srs_task.output else 'Chưa có SRS.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Service_Level_Agreement_Template.md' mô tả chi tiết các cam kết SLA.",
        agent=requirement_analyst_agent, # Changed to requirement_agent as per table
        context=[srs_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành SLA Task ---"),
            write_output(os.path.join(phase_output_dir, "Service_Level_Agreement_Template.md"), str(output)),
            shared_memory.set("phase_2_requirements", "sla_document", str(output)),
            logging.info(f"Đã lưu Service_Level_Agreement_Template.md và cập nhật shared_memory.")
        )
    )

    # Task: Non-Functional Requirements Document (NFRD) (nfr_tasks.py)
    # Inputs: SRS.docx
    # Output: Non_Functional_Requirements.md (already exists, but confirming details)
    nfrd_task = Task(
        description=dedent(f"""
            Dựa trên SRS (có trong context), phát triển Tài liệu Yêu cầu Phi Chức năng (NFRD).
            Bao gồm các yêu cầu về hiệu suất, bảo mật, khả năng mở rộng, khả năng sử dụng, độ tin cậy, v.v.
            --- SRS: {srs_task.output.raw_output if srs_task.output else 'Chưa có SRS.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Non_Functional_Requirements.md' mô tả chi tiết các yêu cầu phi chức năng.",
        agent=requirement_analyst_agent,
        context=[srs_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành NFRD Task ---"),
            write_output(os.path.join(phase_output_dir, "Non_Functional_Requirements.md"), str(output)),
            shared_memory.set("phase_2_requirements", "nfrd_document", str(output)),
            logging.info(f"Đã lưu Non_Functional_Requirements.md và cập nhật shared_memory.")
        )
    )

    # NEW TASK: Privacy and Security Requirements (security_tasks.py)
    # Inputs: Non_Functional_Requirements.md
    # Output: Privacy_and_Security_Requirements.docx (generating as markdown)
    security_requirements_task = Task(
        description=dedent(f"""
            Từ Tài liệu Yêu cầu Phi Chức năng (NFRD) (có trong context),
            chi tiết hóa các yêu cầu về quyền riêng tư và bảo mật.
            Tạo tài liệu đặc tả các yêu cầu bảo mật và quyền riêng tư của hệ thống.
            --- NFRD: {nfrd_task.output.raw_output if nfrd_task.output else 'Chưa có NFRD.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Privacy_and_Security_Requirements.md' mô tả chi tiết các yêu cầu bảo mật và quyền riêng tư.",
        agent=requirement_analyst_agent, # Assign to requirement_agent as per table
        context=[nfrd_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Privacy and Security Requirements Task ---"),
            write_output(os.path.join(phase_output_dir, "Privacy_and_Security_Requirements.md"), str(output)),
            shared_memory.set("phase_2_requirements", "privacy_security_requirements", str(output)),
            logging.info(f"Đã lưu Privacy_and_Security_Requirements.md và cập nhật shared_memory.")
        )
    )

    # NEW TASK: Requirements Inspection Checklist (checklist_tasks.py)
    # Inputs: SRS.docx, RTM.xlsx
    # Output: Requirements_Inspection_Checklist.xlsx (generating as markdown)
    inspection_checklist_task = Task(
        description=dedent(f"""
            Tạo một checklist kiểm tra yêu cầu (Requirements Inspection Checklist)
            dựa trên SRS và RTM. Checklist này sẽ được sử dụng để đảm bảo chất lượng của các yêu cầu.
            --- SRS: {srs_task.output.raw_output if srs_task.output else 'Chưa có SRS.'}
            --- RTM: {rtm_task.output.raw_output if rtm_task.output else 'Chưa có RTM.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Requirements_Inspection_Checklist.md' dùng để kiểm tra chất lượng yêu cầu.",
        agent=requirement_analyst_agent, # Assign to requirement_agent as per table
        context=[srs_task, rtm_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Requirements Inspection Checklist Task ---"),
            write_output(os.path.join(phase_output_dir, "Requirements_Inspection_Checklist.md"), str(output)),
            shared_memory.set("phase_2_requirements", "requirements_inspection_checklist", str(output)),
            logging.info(f"Đã lưu Requirements_Inspection_Checklist.md và cập nhật shared_memory.")
        )
    )

    # NEW TASK: Training Plan (training_tasks.py)
    # Inputs: Use_Case_Diagram.png, CONOPS.docx
    # Output: Training_Plan.docx (generating as markdown)
    training_plan_task = Task(
        description=dedent(f"""
            Xây dựng kế hoạch đào tạo (Training Plan) dựa trên Use Case Diagram và CONOPS.
            Kế hoạch này sẽ bao gồm các nội dung đào tạo cần thiết cho người dùng cuối và các bên liên quan khác.
            --- Use Case Diagram: {shared_memory.get("phase_2_requirements", "use_case_diagram_path") or 'Chưa có Use Case Diagram.'}
            --- CONOPS: {conops_content}
        """),
        expected_output="Tài liệu tiếng Việt 'Training_Plan.md' chi tiết về kế hoạch đào tạo.",
        agent=requirement_analyst_agent, # Assign to requirement_agent as per table
        # We need to ensure use_case_diagram_path is set by the use_case_task callback.
        # CONOPS content is passed directly.
        context=[use_case_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Training Plan Task ---"),
            write_output(os.path.join(phase_output_dir, "Training_Plan.md"), str(output)),
            shared_memory.set("phase_2_requirements", "training_plan", str(output)),
            logging.info(f"Đã lưu Training_Plan.md và cập nhật shared_memory.")
        )
    )

    # Existing validation and research tasks, updated context for better reflection
    # Task: SRS Validation - Now relies on SRS being generated.
    srs_validation_task = Task(
        description=dedent(f"""
            Kiểm tra SRS về tính hoàn chỉnh, rõ ràng, nhất quán và khả thi.
            Phát hiện các mâu thuẫn, thiếu sót hoặc điểm không rõ ràng trong đặc tả kỹ thuật.
            Tạo báo cáo tóm tắt kết quả kiểm tra.
            --- SRS: {srs_task.output.raw_output if srs_task.output else 'Chưa có SRS.'}
        """),
        expected_output="Tài liệu tiếng Việt 'SRS_Validation_Report.md' tóm tắt các phát hiện và đề xuất.",
        agent=project_manager_agent,
        context=[srs_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành SRS Validation Task ---"),
            write_output(os.path.join(phase_output_dir, "SRS_Validation_Report.md"), str(output)),
            shared_memory.set("phase_2_requirements", "srs_validation_report", str(output)),
            logging.info(f"Đã lưu SRS_Validation_Report.md và cập nhật shared_memory.")
        )
    )

    # Task: Requirements Validation (Project Manager's Quality Gate) - Updated context
    requirements_validation_task = Task(
        description=dedent(f"""
            Đánh giá kỹ lưỡng tất cả các tài liệu yêu cầu được tạo ra trong giai đoạn này
            (Project Scope Checklist, BRD, BRD Presentation, SRS, FRD, NFRD, Use Cases & User Stories, RTM,
            Change Impact Analysis, SLA Template, Privacy & Security Requirements, Requirements Inspection Checklist, Training Plan).
            Kiểm tra tính hoàn chỉnh, rõ ràng, nhất quán và khả thi của toàn bộ bộ tài liệu.
            Tạo báo cáo 'Validation_Report_Phase_2.md' tóm tắt kết quả đánh giá,
            liệt kê các điểm cần cải thiện nếu có và xác nhận hoàn thành giai đoạn.
        """),
        expected_output="Tài liệu tiếng Việt 'Validation_Report_Phase_2.md' tóm tắt kết quả đánh giá giai đoạn yêu cầu.",
        agent=project_manager_agent,
        context=[
            scope_task,
            brd_task,
            brd_presentation_task,
            srs_task,
            srs_validation_task,
            frd_task,
            nfrd_task,
            use_case_task, # This produces a docx and sets use_case_diagram_path, user_stories_path
            rtm_task,
            impact_analysis_task, # New
            sla_task,
            security_requirements_task, # New
            inspection_checklist_task, # New
            training_plan_task # New
        ],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Requirements Validation Task ---"),
            write_output(os.path.join(phase_output_dir, "Validation_Report_Phase_2.md"), str(output)),
            shared_memory.set("phase_2_requirements", "validation_report", str(output)),
            logging.info(f"Đã lưu Validation_Report_Phase_2.md và cập nhật shared_memory.")
        )
    )

    # Task: Research Requirements Best Practices (Researcher)
    research_requirements_best_practices = Task(
        description=dedent(f"""
            Nghiên cứu các phương pháp hay nhất (best practices) và tiêu chuẩn ngành trong quản lý yêu cầu phần mềm
            (ví dụ: IEEE, BABOK, MoSCoW, SMART requirements). Tổng hợp kiến thức hỗ trợ các agent khác.
            --- Contextual Documents: {srs_task.output.raw_output if srs_task.output else 'Chưa có SRS.'},
            {brd_task.output.raw_output if brd_task.output else 'Chưa có BRD.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Requirements_Research_Summary.md' tổng hợp các kiến thức nghiên cứu hữu ích.",
        agent=researcher_agent,
        context=[srs_task, brd_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Research Requirements Best Practices Task ---"),
            write_output(os.path.join(phase_output_dir, "Requirements_Research_Summary.md"), str(output)),
            shared_memory.set("phase_2_requirements", "research_summary", str(output)),
            logging.info(f"Đã lưu Requirements_Research_Summary.md và cập nhật shared_memory.")
        )
    )

    return [
        scope_task,
        brd_task,
        brd_presentation_task,
        srs_task,
        frd_task,
        use_case_task,
        rtm_task,
        impact_analysis_task,  # New
        sla_task,
        nfrd_task,
        security_requirements_task, # New
        inspection_checklist_task,  # New
        training_plan_task, # New
        srs_validation_task, # Reordered for better flow but context depends on srs_task
        requirements_validation_task, # Final validation
        research_requirements_best_practices # Research task can run in parallel
    ]