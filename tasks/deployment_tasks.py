import os
import logging
import re
from crewai import Task
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches
import graphviz

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def process_and_create_deployment_doc(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho Kế hoạch Triển khai ---")
    
    dot_blocks = re.findall(r'```dot\s*([\s\S]*?)\s*```', task_output)
    text_content = re.sub(r'```dot\s*[\s\S]*?```', '', task_output).strip()

    deployment_plan_text = text_content # Phần văn bản của kế hoạch triển khai
    deployment_diagram_dot_code = ""

    if len(dot_blocks) >= 1:
        deployment_diagram_dot_code = dot_blocks[0]
        logging.info("Đã trích xuất mã DOT cho Deployment Diagram.")

    # Đảm bảo thư mục con tồn tại cho output của phase này (Phase 6 - Deployment)
    phase_output_dir = os.path.join(output_base_dir_param, "6_deployment")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 6: {phase_output_dir}")

    # --- 1. Tạo Deployment_Plan.md ---
    deployment_plan_md_path = os.path.join(phase_output_dir, "Deployment_Plan.md")
    write_output(deployment_plan_md_path, deployment_plan_text)
    logging.info(f"Đã lưu Deployment_Plan.md vào {deployment_plan_md_path}")
    shared_memory.set("phase_6_deployment", "deployment_plan_md_path", deployment_plan_md_path)

    def write_docx(file_path, content, heading):
        """Tạo file .docx hợp lệ bằng python-docx."""
        try:
            doc = Document()
            doc.add_heading(heading, level=1)
            doc.add_paragraph(content)
            doc.save(file_path)
            logging.info(f"Đã lưu file DOCX hợp lệ: {file_path}")
        except Exception as e:
            logging.error(f"Lỗi khi lưu file DOCX {file_path}: {str(e)}")


        production_impl_plan_docx_path = os.path.join(phase_output_dir, "Production_Implementation_Plan.docx")
        doc.save(production_impl_plan_docx_path)
        logging.info(f"Đã lưu tài liệu Production_Implementation_Plan.docx vào {production_impl_plan_docx_path}")
        shared_memory.set("phase_6_deployment", "production_implementation_plan_docx_path", production_impl_plan_docx_path)

        logging.info(f"--- Hoàn thành xử lý output cho Kế hoạch Triển khai ---")

def process_turnover_documents(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho các tài liệu Bàn giao ---")

    phase_output_dir = os.path.join(output_base_dir_param, "6_deployment")
    os.makedirs(phase_output_dir, exist_ok=True) # Đảm bảo thư mục tồn tại

    # Sử dụng regex để tìm các khối nội dung được đánh dấu
    form_content = re.search(r'```form\s*([\s\S]*?)\s*```', task_output)
    install_content = re.search(r'```install_guide\s*([\s\S]*?)\s*```', task_output)
    ops_content = re.search(r'```ops_guide\s*([\s\S]*?)\s*```', task_output)

    # Production_Turnover_Approval_Form.docx
    form_text = form_content.group(1).strip() if form_content else "Không tìm thấy nội dung mẫu phê duyệt."
    form_doc_path = os.path.join(phase_output_dir, "Production_Turnover_Approval_Form.docx")
    write_output(form_doc_path, form_text)
    shared_memory.set("phase_6_deployment", "production_turnover_approval_form_path", form_doc_path)
    logging.info(f"Đã lưu Production_Turnover_Approval_Form.docx và cập nhật shared_memory.")

    # Installation_Guide.docx
    install_text = install_content.group(1).strip() if install_content else "Không tìm thấy nội dung hướng dẫn cài đặt."
    install_doc_path = os.path.join(phase_output_dir, "Installation_Guide.docx")
    write_output(install_doc_path, install_text)
    shared_memory.set("phase_6_deployment", "installation_guide_path", install_doc_path)
    logging.info(f"Đã lưu Installation_Guide.docx và cập nhật shared_memory.")

    # Operations_Guide.docx
    ops_text = ops_content.group(1).strip() if ops_content else "Không tìm thấy nội dung hướng dẫn vận hành."
    ops_doc_path = os.path.join(phase_output_dir, "Operations_Guide.docx")
    write_output(ops_doc_path, ops_text)
    shared_memory.set("phase_6_deployment", "operations_guide_path", ops_doc_path)
    logging.info(f"Đã lưu Operations_Guide.docx và cập nhật shared_memory.")

    logging.info(f"--- Hoàn thành xử lý output cho các tài liệu Bàn giao ---")


def create_deployment_tasks(devops_engineer_agent, project_manager_agent, researcher_agent, output_base_dir):
    """
    Tạo các task liên quan đến giai đoạn Triển khai, Bàn giao và Giám sát.
    Args:
        devops_engineer_agent: Agent chính cho Deployment.
        project_manager_agent, researcher_agent: Các agent chung.
        output_base_dir: Đường dẫn thư mục base.
    Returns:
        list: Danh sách các Task đã tạo.
    """
    # Lấy dữ liệu từ các phase trước
    test_execution_report = shared_memory.get("phase_5_testing", "test_execution_report")
    build_deploy_process = shared_memory.get("phase_4_development", "build_deploy_process")
    system_architecture = shared_memory.get("phase_3_design", "system_architecture_document")

    deployment_context = f"Test Execution Report (Snippet): {test_execution_report[:500] if test_execution_report else 'N/A'}\n" \
                         f"Build/Deployment Process (Snippet): {build_deploy_process[:500] if build_deploy_process else 'N/A'}\n" \
                         f"System Architecture (Snippet): {system_architecture[:500] if system_architecture else 'N/A'}"
    
    if not any([test_execution_report, build_deploy_process, system_architecture]):
        logging.warning("Previous phase documents missing for deployment tasks.")
        deployment_context = "Không có tài liệu liên quan nào được tìm thấy."

    # Tạo thư mục con cho Phase 6 Deployment (nếu chưa có)
    phase_output_dir = os.path.join(output_base_dir, "6_deployment")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 6: {phase_output_dir}")

    # --- 1. Task: Deployment Plan ---
    # Nhiệm vụ tạo Deployment_Plan.md và Production_Implementation_Plan.docx
    deployment_plan_task = Task(
        description=(
            f"Dựa trên các tài liệu thiết kế (kiến trúc hệ thống), kết quả kiểm thử và quy trình build/deploy, "
            f"tạo Kế hoạch Triển khai chi tiết. Output phải bao gồm:\n"
            f"1. Nội dung văn bản của kế hoạch triển khai (ví dụ: các bước, yêu cầu môi trường, rollback plan).\n"
            f"2. Mã nguồn Graphviz DOT cho một Deployment Diagram minh họa cấu trúc vật lý của hệ thống "
            f"và nơi các thành phần phần mềm sẽ được triển khai. "
            f"Cấu trúc đầu ra của bạn phải là phần văn bản kế hoạch, sau đó là mã DOT (trong '```dot\\n...\\n```').\n"
            f"--- Context: {deployment_context}"
        ),
        expected_output=(
            "Một chuỗi văn bản (string) bao gồm:\n"
            "1. Phần mô tả Kế hoạch Triển khai (để lưu vào Deployment_Plan.md).\n"
            "2. Tiếp theo là mã Graphviz DOT cho Deployment Diagram được bọc trong '```dot\\n...\\n```' (để chèn vào Production_Implementation_Plan.docx).\n"
            "Đảm bảo mã DOT đúng cú pháp để có thể render thành hình ảnh."
        ),
        agent=devops_engineer_agent,
        callback=lambda output: process_and_create_deployment_doc(str(output), output_base_dir)
    )

    # --- 2. Task: Handover Documents (bao gồm Approval Form, Installation Guide, Operations Guide) ---
    handover_documents_task = Task( 
        description=(
            f"Tạo ba tài liệu sau đây trong một phản hồi duy nhất, sử dụng các khối mã riêng biệt:\n"
            f"1. **Production Turnover Approval Form.docx**: Mẫu phê duyệt bàn giao sản xuất, xác nhận việc bàn giao thành công. "
            f"Bao gồm thông tin dự án, danh sách các tài liệu bàn giao, các hạng mục kiểm tra, và mục chữ ký.\n"
            f"2. **Installation_Guide.docx**: Hướng dẫn cài đặt chi tiết từng bước cho hệ thống trên môi trường sản xuất, "
            f"bao gồm yêu cầu, phụ thuộc và các lệnh cần thiết.\n"
            f"3. **Operations_Guide.docx**: Hướng dẫn vận hành hàng ngày cho hệ thống, bao gồm quy trình vận hành, "
            f"khắc phục sự cố phổ biến, sao lưu/phục hồi và quản lý hiệu suất.\n"
            f"Cấu trúc đầu ra của bạn phải như sau:\n"
            f"```form\n[Nội dung Mẫu Phê duyệt Bàn giao]\n```\n"
            f"```install_guide\n[Nội dung Hướng dẫn Cài đặt]\n```\n"
            f"```ops_guide\n[Nội dung Hướng dẫn Vận hành]\n```\n"
            f"--- Production Implementation Plan: {shared_memory.get('phase_6_deployment', 'production_implementation_plan_docx_path')}"
        ),
        expected_output=(
            "Một chuỗi văn bản (string) chứa nội dung cho ba tài liệu:\n"
            "1. Nội dung cho 'Production_Turnover_Approval_Form.docx' trong khối '```form...```'.\n"
            "2. Nội dung cho 'Installation_Guide.docx' trong khối '```install_guide...```'.\n"
            "3. Nội dung cho 'Operations_Guide.docx' trong khối '```ops_guide...```'.\n"
            "Đảm bảo mỗi khối mã chứa nội dung đầy đủ và phù hợp cho từng tài liệu."
        ),
        agent=devops_engineer_agent,
        context=[deployment_plan_task], 
        callback=lambda output: process_turnover_documents(str(output), output_base_dir)
    )

    # --- 3. Task: Monitoring Setup Guide ---
    monitoring_setup_guide_task = Task(
        description=(
            f"Phát triển Hướng dẫn Thiết lập Giám sát và Cảnh báo (Monitoring and Alerting Setup Guide) cho hệ thống đã triển khai. "
            f"Xác định các chỉ số quan trọng (KPIs), công cụ giám sát, ngưỡng cảnh báo, và quy trình xử lý cảnh báo. "
            f"Tài liệu này nên tập trung vào các bước cấu hình cụ thể để thiết lập hệ thống giám sát.\n"
            f"--- Production Implementation Plan: {shared_memory.get('phase_6_deployment', 'production_implementation_plan_docx_path')}"
        ),
        expected_output="Tài liệu tiếng Việt 'Monitoring_and_Alerting_Setup_Guide.md' mô tả các bước thiết lập giám sát.",
        agent=devops_engineer_agent,
        context=[deployment_plan_task], 
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Monitoring and Alerting Setup Guide Task ---"),
            write_output(os.path.join(phase_output_dir, "Monitoring_and_Alerting_Setup_Guide.md"), str(output)),
            shared_memory.set("phase_6_deployment", "monitoring_and_alerting_setup_guide_path", str(output)),
            logging.info(f"Đã lưu Monitoring_and_Alerting_Setup_Guide.md và cập nhật shared_memory.")
        )
    )

    # --- 4. Task: Deployment Quality Gate (Project Manager) ---
    deployment_validation_task = Task(
        description=(
            f"Đánh giá kỹ lưỡng các tài liệu triển khai, bàn giao và giám sát "
            f"(Deployment Plan, Production Implementation Plan, Production Turnover Approval Form, Installation Guide, Operations Guide, Monitoring and Alerting Setup Guide). "
            f"Kiểm tra tính hoàn chỉnh, khả thi, và rủi ro. "
            f"Tạo báo cáo 'Validation_Report_Phase_6.md' tóm tắt kết quả đánh giá, "
            f"liệt kê các điểm cần cải thiện nếu có và xác nhận hoàn thành giai đoạn."
        ),
        expected_output="Tài liệu tiếng Việt 'Validation_Report_Phase_6.md' tóm tắt kết quả đánh giá giai đoạn triển khai.",
        agent=project_manager_agent,
        context=[
            deployment_plan_task, 
            handover_documents_task, # Task này giờ bao gồm cả Installation và Operations Guide
            monitoring_setup_guide_task
        ],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Deployment Validation Task ---"),
            write_output(os.path.join(phase_output_dir, "Validation_Report_Phase_6.md"), str(output)),
            shared_memory.set("phase_6_deployment", "validation_report_path", str(output)),
            logging.info(f"Đã lưu Validation_Report_Phase_6.md và cập nhật shared_memory.")
        )
    )

    # --- 5. Task: Research Deployment Best Practices (Researcher) ---
    research_deployment_best_practices = Task(
        description=(
            f"Nghiên cứu các phương pháp hay nhất (best practices) trong triển khai và vận hành phần mềm "
            f"(ví dụ: CI/CD, Infrastructure as Code, Blue/Green Deployment, Canary Releases). "
            f"Tổng hợp kiến thức hỗ trợ các agent khác.\n"
            f"--- Production Implementation Plan: {shared_memory.get('phase_6_deployment', 'production_implementation_plan_docx_path')}"
        ),
        expected_output="Tài liệu tiếng Việt 'Deployment_Research_Summary.md' tổng hợp các kiến thức nghiên cứu hữu ích.",
        agent=researcher_agent,
        context=[deployment_plan_task], 
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Research Deployment Best Practices Task ---"),
            write_output(os.path.join(phase_output_dir, "Deployment_Research_Summary.md"), str(output)),
            shared_memory.set("phase_6_deployment", "research_summary_path", str(output)),
            logging.info(f"Đã lưu Deployment_Research_Summary.md và cập nhật shared_memory.")
        )
    )

    return [
        deployment_plan_task,
        handover_documents_task,
        monitoring_setup_guide_task,
        deployment_validation_task,
        research_deployment_best_practices
    ]