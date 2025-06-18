import os
import logging
import re
from crewai import Task
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches
from textwrap import dedent
import graphviz
import pandas as pd

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def read_file_content(file_path):
    """
    Đọc nội dung từ file .docx, .xlsx hoặc file văn bản (.txt, .md).
    Args:
        file_path (str): Đường dẫn tới file.
    Returns:
        str: Nội dung của file hoặc thông báo lỗi nếu không đọc được.
    """
    if not file_path or not os.path.exists(file_path):
        logging.warning(f"File không tồn tại hoặc đường dẫn rỗng: {file_path}")
        return "Nội dung tài liệu không tìm thấy."

    try:
        if file_path.endswith('.docx'):
            logging.info(f"Đọc file .docx: {file_path}")
            doc = Document(file_path)
            content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            return content if content.strip() else "Nội dung file .docx rỗng."
        elif file_path.endswith('.xlsx'):
            logging.info(f"Đọc file .xlsx: {file_path}")
            df = pd.read_excel(file_path)
            return df.to_string() if not df.empty else "Nội dung file .xlsx rỗng."
        else:
            logging.info(f"Đọc file văn bản: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        logging.error(f"Lỗi khi đọc file {file_path}: {e}", exc_info=True)
        return f"Lỗi khi đọc file: {str(e)}"

def clean_text_for_docx(text: str) -> str:
    """
    Loại bỏ các ký tự có thể gây lỗi khi ghi vào file .docx hoặc làm phẳng Markdown cơ bản.
    """
    clean_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text).strip()
    clean_text = re.sub(r'#{1,6}\s*', '', clean_text)
    clean_text = clean_text.replace('**', '').replace('__', '')
    clean_text = clean_text.replace('*', '').replace('_', '')
    clean_text = re.sub(r'^- ', '', clean_text, flags=re.MULTILINE)
    clean_text = re.sub(r'^\d+\.\s*', '', clean_text, flags=re.MULTILINE)
    return clean_text

def process_and_save_docx(task_output: str, file_path: str, document_title: str):
    """
    Tạo tài liệu Word (.docx) từ output của agent.
    Cố gắng phân tích các tiêu đề và đoạn văn cơ bản.
    """
    logging.info(f"--- Bắt đầu xử lý và lưu tài liệu Word: {file_path} ---")
    doc = Document()
    doc.add_heading(document_title, level=0)

    cleaned_output = clean_text_for_docx(task_output)
    lines = cleaned_output.split('\n')
    current_section = []

    for line in lines:
        stripped_line = line.strip()
        if stripped_line:
            current_section.append(stripped_line)
        else:
            if current_section:
                doc.add_paragraph(' '.join(current_section))
                current_section = []
    if current_section:
        doc.add_paragraph(' '.join(current_section))

    try:
        doc.save(file_path)
        logging.info(f"Đã lưu thành công tài liệu Word: {file_path}")
    except Exception as e:
        logging.error(f"Lỗi khi lưu tài liệu Word {file_path}: {e}", exc_info=True)
        write_output(file_path + ".error_raw.txt", task_output)

def process_and_create_planning_doc(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho Kế hoạch Dự án ---")
    
    phase_output_dir = os.path.join(output_base_dir_param, "1_planning")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 1: {phase_output_dir}")

    # Tách các khối mã DOT (giả định WBS là một sơ đồ cây)
    dot_blocks = re.findall(r'```dot\s*([\s\S]*?)\s*```', task_output)
    
    # Chia output thành các phần dựa trên tiêu đề Markdown
    wbs_doc_match = re.search(r'### WBS Document\n(.*?)(?=\n### WBS Dictionary|\Z)', task_output, re.DOTALL)
    wbs_dict_match = re.search(r'### WBS Dictionary\n(.*?)(?=\n### WBS Resource Template|\Z)', task_output, re.DOTALL)
    wbs_resource_template_match = re.search(r'### WBS Resource Template\n(.*?)(?=\n### Project Plan XML|\Z)', task_output, re.DOTALL)
    project_plan_xml_match = re.search(r'### Project Plan XML\n(.*?)(?=\Z)', task_output, re.DOTALL)

    wbs_doc_content = wbs_doc_match.group(1).strip() if wbs_doc_match else "Không có nội dung WBS Document."
    wbs_dict_content = wbs_dict_match.group(1).strip() if wbs_dict_match else "Không có nội dung WBS Dictionary."
    wbs_resource_template_content = wbs_resource_template_match.group(1).strip() if wbs_resource_template_match else "Không có nội dung WBS Resource Template."
    project_plan_xml_content = project_plan_xml_match.group(1).strip() if project_plan_xml_match else "Không có nội dung Project Plan XML."

    # Tạo WBS Document (docx)
    process_and_save_docx(wbs_doc_content, os.path.join(phase_output_dir, "WBS.docx"), "Cấu trúc Phân chia Công việc (WBS)")
    shared_memory.set("phase_1_planning", "wbs_document_path", os.path.join(phase_output_dir, "WBS.docx"))
    logging.info(f"Đã lưu WBS.docx và cập nhật shared_memory.")

    # Tạo WBS Dictionary (docx)
    process_and_save_docx(wbs_dict_content, os.path.join(phase_output_dir, "WBS_Dictionary.docx"), "Từ điển WBS")
    shared_memory.set("phase_1_planning", "wbs_dictionary_path", os.path.join(phase_output_dir, "WBS_Dictionary.docx"))
    logging.info(f"Đã lưu WBS_Dictionary.docx và cập nhật shared_memory.")

    # Tạo WBS Resource Template (docx - mô phỏng xlsx)
    process_and_save_docx(wbs_resource_template_content, os.path.join(phase_output_dir, "WBS_Resource_Template.docx"), "Mẫu Tài nguyên WBS")
    shared_memory.set("phase_1_planning", "wbs_resource_template_path", os.path.join(phase_output_dir, "WBS_Resource_Template.docx"))
    logging.info(f"Đã lưu WBS_Resource_Template.docx và cập nhật shared_memory.")

    # Tạo Project Plan XML (txt - mô phỏng xml)
    write_output(os.path.join(phase_output_dir, "Project_Plan.txt"), project_plan_xml_content)
    shared_memory.set("phase_1_planning", "project_plan_xml_path", os.path.join(phase_output_dir, "Project_Plan.txt"))
    logging.info(f"Đã lưu Project_Plan.txt và cập nhật shared_memory.")

    # Render WBS Diagram (nếu có mã DOT)
    if len(dot_blocks) >= 1:
        wbs_dot_code = dot_blocks[0]
        try:
            graph_wbs = graphviz.Source(wbs_dot_code, format='png', engine='dot')
            wbs_img_path = os.path.join(phase_output_dir, "WBS_Diagram.png")
            graph_wbs.render(wbs_img_path.rsplit('.', 1)[0], view=False, cleanup=True)
            shared_memory.set("phase_1_planning", "wbs_diagram_path", wbs_img_path)
            logging.info(f"Đã tạo WBS Diagram: {wbs_img_path}")
        except Exception as e:
            logging.error(f"Lỗi khi tạo WBS Diagram: {e}", exc_info=True)
    else:
        logging.warning("Agent không tạo ra mã DOT cho WBS Diagram.")

    logging.info(f"--- Hoàn thành xử lý output cho Kế hoạch Dự án ---")

def create_planning_tasks(planning_orchestrator_agent, project_manager_agent, researcher_agent, output_base_dir):
    """
    Tạo các task liên quan đến lập kế hoạch dự án.
    Args:
        planning_orchestrator_agent: Agent chính cho Planning.
        project_manager_agent, researcher_agent: Các agent chung.
        output_base_dir: Đường dẫn thư mục base.
    Returns:
        list: Danh sách các Task đã tạo.
    """
    # Lấy outputs từ Phase 0 Initiation
    project_charter_content = shared_memory.get("phase_0_initiation", "project_charter_path")
    business_case_content = shared_memory.get("phase_0_initiation", "business_case_path")
    feasibility_report_content = shared_memory.get("phase_0_initiation", "feasibility_report_path")
    conops_document_content = shared_memory.get("phase_0_initiation", "concept_of_operations_path")
    initiate_project_checklist_content = shared_memory.get("phase_0_initiation", "initiate_project_checklist_path")
    preliminary_schedule_content = shared_memory.get("phase_0_initiation", "preliminary_schedule_path")
    project_resource_plan_content = shared_memory.get("phase_0_initiation", "project_resource_plan_path")
    project_team_definition_content = shared_memory.get("phase_0_initiation", "project_team_definition_path")
    risk_assessment_document_content = shared_memory.get("phase_0_initiation", "risk_assessment_document_path")

    # Tạo thư mục con cho Phase 1 Planning
    phase_output_dir = os.path.join(output_base_dir, "1_planning")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 1: {phase_output_dir}")

    # Task: Project Plan Creation
    project_plan_task = Task(
        description=dedent(f"""
            Dựa trên Project Charter sau, phát triển một Project Plan (Kế hoạch Dự án) chi tiết.
            Bao gồm các phần chính như: Mục tiêu chi tiết, Phạm vi dự án, Lịch trình (Timeline),
            Nguồn lực (Resources), Ngân sách sơ bộ, Quản lý rủi ro, và Kế hoạch truyền thông.
            --- Project Charter: {read_file_content(project_charter_content)}
        """),
        expected_output="Tài liệu tiếng Việt 'Project_Plan.docx' đầy đủ và có cấu trúc.",
        agent=planning_orchestrator_agent,
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Project Plan Task ---"),
            process_and_save_docx(str(output), os.path.join(phase_output_dir, "Project_Plan.docx"), "Kế hoạch Dự án"),
            shared_memory.set("phase_1_planning", "project_plan_path", os.path.join(phase_output_dir, "Project_Plan.docx")),
            logging.info(f"Đã lưu Project_Plan.docx và cập nhật shared_memory.")
        )
    )

    # Task: PMO Tasks
    pmo_tasks = Task(
        description=dedent(f"""
            Sử dụng 'Initiate_Project_Checklist' để tạo các checklist chất lượng:
            1. 'PMO_Checklist.xlsx' (mô phỏng .md)
            2. 'COBIT_Checklist.xlsx' (mô phỏng .md)
            3. 'QA_Checklist.xlsx' (mô phỏng .md)
            Sử dụng các tiêu đề rõ ràng (ví dụ: "### PMO CHECKLIST") để phân tách các phần trong output.
        """),
        expected_output="""Một chuỗi văn bản (string) chứa nội dung của các checklist PMO, COBIT, QA được phân tách rõ ràng.
        Các checklist này phải chi tiết và bám sát nội dung của 'Initiate_Project_Checklist'.""",
        agent=planning_orchestrator_agent,
        context=[project_plan_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành PMO Tasks ---"),
            write_output(
                os.path.join(phase_output_dir, "PMO_Checklist.md"),
                re.search(r'### PMO CHECKLIST\n(.*?)(?=\n### COBIT CHECKLIST|\Z)', output.raw, re.DOTALL).group(1).strip()
                if re.search(r'### PMO CHECKLIST\n(.*?)(?=\n### COBIT CHECKLIST|\Z)', output.raw, re.DOTALL)
                else "Không có PMO Checklist."
            ),
            write_output(
                os.path.join(phase_output_dir, "COBIT_Checklist.md"),
                re.search(r'### COBIT CHECKLIST\n(.*?)(?=\n### QA CHECKLIST|\Z)', output.raw, re.DOTALL).group(1).strip()
                if re.search(r'### COBIT CHECKLIST\n(.*?)(?=\n### QA CHECKLIST|\Z)', output.raw, re.DOTALL)
                else "Không có COBIT Checklist."
            ),
            write_output(
                os.path.join(phase_output_dir, "QA_Checklist.md"),
                re.search(r'### QA CHECKLIST\n(.*?)(?=\Z)', output.raw, re.DOTALL).group(1).strip()
                if re.search(r'### QA CHECKLIST\n(.*?)(?=\Z)', output.raw, re.DOTALL)
                else "Không có QA Checklist."
            ),
            shared_memory.set("phase_1_planning", "pmo_checklist_path", os.path.join(phase_output_dir, "PMO_Checklist.md")),
            shared_memory.set("phase_1_planning", "cobit_checklist_path", os.path.join(phase_output_dir, "COBIT_Checklist.md")),
            shared_memory.set("phase_1_planning", "qa_checklist_path", os.path.join(phase_output_dir, "QA_Checklist.md")),
            logging.info(f"Đã lưu các Checklist và cập nhật shared_memory.")
        )
    )

    # Task: Statement of Work (SOW)
    sow_tasks = Task(
        description=dedent(f"""
            Tạo 'Statement_of_Work.docx' dựa trên 'Project_Charter.docx' và 'CONOPS.docx'.
            SOW phải mô tả chi tiết phạm vi công việc, sản phẩm bàn giao, thời gian biểu và các điều khoản khác.
            --- Project Charter: {read_file_content(project_charter_content)}
            --- CONOPS: {read_file_content(conops_document_content)}
        """),
        expected_output="Tài liệu tiếng Việt 'Statement_of_Work.docx' chi tiết.",
        agent=planning_orchestrator_agent,
        context=[project_plan_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành SOW Tasks ---"),
            process_and_save_docx(str(output), os.path.join(phase_output_dir, "Statement_of_Work.docx"), "Tuyên bố Công việc"),
            shared_memory.set("phase_1_planning", "sow_document_path", os.path.join(phase_output_dir, "Statement_of_Work.docx")),
            logging.info(f"Đã lưu Statement_of_Work.docx và cập nhật shared_memory.")
        )
    )

    # Task: Approval Tasks
    approval_tasks = Task(
        description=dedent(f"""
            Tạo 'Project_Approval_Document.docx' dựa trên 'Statement_of_Work.docx' và 'Feasibility_Report.docx'.
            Tài liệu phê duyệt cần nêu rõ các điều kiện, chữ ký phê duyệt và các thông tin liên quan khác.
            --- SOW: {read_file_content(shared_memory.get("phase_1_planning", "sow_document_path"))}
            --- Feasibility Report: {read_file_content(feasibility_report_content)}
        """),
        expected_output="Tài liệu tiếng Việt 'Project_Approval_Document.docx' chính thức.",
        agent=planning_orchestrator_agent,
        context=[sow_tasks],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Approval Tasks ---"),
            process_and_save_docx(str(output), os.path.join(phase_output_dir, "Project_Approval_Document.docx"), "Tài liệu Phê duyệt Dự án"),
            shared_memory.set("phase_1_planning", "project_approval_document_path", os.path.join(phase_output_dir, "Project_Approval_Document.docx")),
            logging.info(f"Đã lưu Project_Approval_Document.docx và cập nhật shared_memory.")
        )
    )

    # Task: Costing Tasks
    costing_tasks = Task(
        description=dedent(f"""
            Dựa trên 'Preliminary_Schedule.xlsx' và 'Project_Resource_Plan.xlsx', tạo các tài liệu ước tính chi phí:
            1. 'Cost_Estimation_Worksheet.xlsx' (mô phỏng .md)
            2. 'Development_Estimation.xlsx' (mô phỏng .md)
            3. 'Capex_Opex_Comparison.xlsx' (mô phỏng .md)
            Sử dụng các tiêu đề rõ ràng (ví dụ: "### COST ESTIMATION WORKSHEET") để phân tách các phần trong output.
            --- Preliminary Schedule: {read_file_content(preliminary_schedule_content)}
            --- Project Resource Plan: {read_file_content(project_resource_plan_content)}
        """),
        expected_output="""Một chuỗi văn bản (string) chứa nội dung của các tài liệu ước tính chi phí được phân tách rõ ràng.""",
        agent=planning_orchestrator_agent,
        context=[project_plan_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Costing Tasks ---"),
            write_output(os.path.join(phase_output_dir, "Cost_Estimation_Worksheet.md"), 
                        re.search(r'### COST ESTIMATION WORKSHEET\n(.*?)(?=\n### DEVELOPMENT ESTIMATION|\Z)', output.raw, re.DOTALL).group(1).strip() 
                        if re.search(r'### COST ESTIMATION WORKSHEET\n(.*?)(?=\n### DEVELOPMENT ESTIMATION|\Z)', output.raw, re.DOTALL) 
                        else "Không có Cost Estimation Worksheet."),
            write_output(os.path.join(phase_output_dir, "Development_Estimation.md"), 
                        re.search(r'### DEVELOPMENT ESTIMATION\n(.*?)(?=\n### CAPEX OPEX COMPARISON|\Z)', output.raw, re.DOTALL).group(1).strip() 
                        if re.search(r'### DEVELOPMENT ESTIMATION\n(.*?)(?=\n### CAPEX OPEX COMPARISON|\Z)', output.raw, re.DOTALL) 
                        else "Không có Development Estimation."),
            write_output(os.path.join(phase_output_dir, "Capex_Opex_Comparison.md"), 
                        re.search(r'### CAPEX OPEX COMPARISON\n(.*?)(?=\Z)', output.raw, re.DOTALL).group(1).strip() 
                        if re.search(r'### CAPEX OPEX COMPARISON\n(.*?)(?=\Z)', output.raw, re.DOTALL) 
                        else "Không có Capex Opex Comparison."),
            shared_memory.set("phase_1_planning", "cost_estimation_worksheet_path", os.path.join(phase_output_dir, "Cost_Estimation_Worksheet.md")),
            shared_memory.set("phase_1_planning", "development_estimation_path", os.path.join(phase_output_dir, "Development_Estimation.md")),
            shared_memory.set("phase_1_planning", "capex_opex_comparison_path", os.path.join(phase_output_dir, "Capex_Opex_Comparison.md")),
            logging.info(f"Đã lưu các tài liệu ước tính chi phí và cập nhật shared_memory.")
        )
    )

    # Task: Org Chart Tasks
    org_chart_tasks = Task(
        description=dedent(f"""
            Dựa trên 'Project_Team_Definition.docx', tạo:
            1. 'Org_Chart.png' (mô phỏng .md hoặc mô tả văn bản)
            2. 'Roles_Responsibilities_Matrix.xlsx' (mô phỏng .md)
            3. 'Approvals_Matrix.xlsx' (mô phỏng .md)
            Sử dụng các tiêu đề rõ ràng (ví dụ: "### ORGANIZATION CHART") để phân tách các phần trong output.
            --- Project Team Definition: {read_file_content(project_team_definition_content)}
        """),
        expected_output="""Một chuỗi văn bản (string) chứa nội dung của sơ đồ tổ chức, ma trận vai trò/trách nhiệm và ma trận phê duyệt được phân tách rõ ràng.""",
        agent=planning_orchestrator_agent,
        context=[project_plan_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Org Chart Tasks ---"),
            write_output(os.path.join(phase_output_dir, "Org_Chart.md"), 
                        re.search(r'### ORGANIZATION CHART\n(.*?)(?=\n### ROLES AND RESPONSIBILITIES MATRIX|\Z)', output.raw, re.DOTALL).group(1).strip() 
                        if re.search(r'### ORGANIZATION CHART\n(.*?)(?=\n### ROLES AND RESPONSIBILITIES MATRIX|\Z)', output.raw, re.DOTALL) 
                        else "Không có Org Chart."),
            write_output(os.path.join(phase_output_dir, "Roles_Responsibilities_Matrix.md"), 
                        re.search(r'### ROLES AND RESPONSIBILITIES MATRIX\n(.*?)(?=\n### APPROVALS MATRIX|\Z)', output.raw, re.DOTALL).group(1).strip() 
                        if re.search(r'### ROLES AND RESPONSIBILITIES MATRIX\n(.*?)(?=\n### APPROVALS MATRIX|\Z)', output.raw, re.DOTALL) 
                        else "Không có Roles/Responsibilities Matrix."),
            write_output(os.path.join(phase_output_dir, "Approvals_Matrix.md"), 
                        re.search(r'### APPROVALS MATRIX\n(.*?)(?=\Z)', output.raw, re.DOTALL).group(1).strip() 
                        if re.search(r'### APPROVALS MATRIX\n(.*?)(?=\Z)', output.raw, re.DOTALL) 
                        else "Không có Approvals Matrix."),
            shared_memory.set("phase_1_planning", "org_chart_path", os.path.join(phase_output_dir, "Org_Chart.md")),
            shared_memory.set("phase_1_planning", "roles_responsibilities_matrix_path", os.path.join(phase_output_dir, "Roles_Responsibilities_Matrix.md")),
            shared_memory.set("phase_1_planning", "approvals_matrix_path", os.path.join(phase_output_dir, "Approvals_Matrix.md")),
            logging.info(f"Đã lưu các tài liệu Tổ chức và cập nhật shared_memory.")
        )
    )

    # Task: Work Breakdown Structure (WBS)
    wbs_task = Task(
        description=dedent(f"""
            Dựa trên 'Statement_of_Work.docx' và 'Project_Approval_Document.docx', chi tiết hóa các công việc thành một Cấu trúc Phân chia Công việc (WBS)
            theo định dạng phân cấp. Xác định các gói công việc chính và các đầu việc con cho từng giai đoạn dự án.
            Bạn phải **tạo mã nguồn Graphviz DOT** để biểu diễn WBS dưới dạng sơ đồ cây phân cấp.
            Cấu trúc đầu ra của bạn phải bao gồm các phần sau, được phân tách bằng các tiêu đề Markdown rõ ràng:
            1. '### WBS Document': Mô tả WBS chính (sẽ được lưu thành .docx).
            2. '### WBS Dictionary': Định nghĩa các mục trong WBS (sẽ được lưu thành .docx).
            3. '### WBS Resource Template': Bảng hoặc mô tả tài nguyên cần thiết cho WBS (sẽ được lưu thành .docx - mô phỏng xlsx).
            4. '### Project Plan XML': Mô phỏng nội dung XML cho Project Plan (sẽ được lưu thành .txt).
            5. Cuối cùng, mã Graphviz DOT cho WBS Diagram (trong '```dot\n...\n```').
            Đảm bảo mã DOT đúng cú pháp để có thể render thành hình ảnh.
            --- SOW: {read_file_content(shared_memory.get("phase_1_planning", "sow_document_path"))}
            --- Project Approval Document: {read_file_content(shared_memory.get("phase_1_planning", "project_approval_document_path"))}
        """),
        expected_output=(
            "Một chuỗi văn bản (string) bao gồm:\n"
            "1. Phần mô tả Cấu trúc Phân chia Công việc (WBS).\n"
            "2. Phần 'WBS Dictionary'.\n"
            "3. Phần 'WBS Resource Template'.\n"
            "4. Phần 'Project Plan XML'.\n"
            "5. Tiếp theo là mã Graphviz DOT cho WBS Diagram được bọc trong '```dot\n...\n```'.\n"
            "Đảm bảo mã DOT đúng cú pháp để có thể render thành hình ảnh."
        ),
        agent=planning_orchestrator_agent,
        context=[sow_tasks, approval_tasks],
        callback=lambda output: process_and_create_planning_doc(str(output), output_base_dir)
    )

    # Task: Risk Management Plan
    risk_plan_task = Task(
        description=dedent(f"""
            Dựa trên 'Risk_Assessment_Document.md', phát triển Kế hoạch Quản lý Rủi ro.
            Xác định các rủi ro tiềm ẩn, phân tích tác động và khả năng xảy ra, và đề xuất các chiến lược giảm thiểu.
            Tạo các tài liệu sau, được phân tách bằng các tiêu đề rõ ràng:
            1. '### Risk Information Form': Biểu mẫu thông tin rủi ro (sẽ được lưu thành .docx).
            2. '### Risk Analysis Plan': Kế hoạch phân tích rủi ro (sẽ được lưu thành .md).
            3. '### Risk Management Plan': Kế hoạch quản lý rủi ro tổng thể (sẽ được lưu thành .docx).
            --- Risk Assessment Document: {read_file_content(risk_assessment_document_content)}
        """),
        expected_output="""Một chuỗi văn bản (string) chứa nội dung của 'Risk_Information_Form.docx', 'Risk_Analysis_Plan.md', và 'Risk_Management_Plan.docx' được phân tách rõ ràng.""",
        agent=planning_orchestrator_agent,
        context=[project_plan_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Risk Management Plan Task ---"),
            process_and_save_docx(
                re.search(r'### RISK INFORMATION FORM\n(.*?)(?=\n### RISK ANALYSIS PLAN|\Z)', output.raw, re.DOTALL).group(1).strip() 
                if re.search(r'### RISK INFORMATION FORM\n(.*?)(?=\n### RISK ANALYSIS PLAN|\Z)', output.raw, re.DOTALL) 
                else "Không có Risk Information Form.", 
                os.path.join(phase_output_dir, "Risk_Information_Form.docx"), 
                "Biểu mẫu Thông tin Rủi ro"
            ),
            write_output(
                os.path.join(phase_output_dir, "Risk_Analysis_Plan.md"), 
                re.search(r'### RISK ANALYSIS PLAN\n(.*?)(?=\n### RISK MANAGEMENT PLAN|\Z)', output.raw, re.DOTALL).group(1).strip() 
                if re.search(r'### RISK ANALYSIS PLAN\n(.*?)(?=\n### RISK MANAGEMENT PLAN|\Z)', output.raw, re.DOTALL) 
                else "Không có Risk Analysis Plan."
            ),
            process_and_save_docx(
                re.search(r'### RISK MANAGEMENT PLAN\n(.*?)(?=\Z)', output.raw, re.DOTALL).group(1).strip() 
                if re.search(r'### RISK MANAGEMENT PLAN\n(.*?)(?=\Z)', output.raw, re.DOTALL) 
                else "Không có Risk Management Plan.", 
                os.path.join(phase_output_dir, "Risk_Management_Plan.docx"), 
                "Kế hoạch Quản lý Rủi ro"
            ),
            shared_memory.set("phase_1_planning", "risk_information_form_path", os.path.join(phase_output_dir, "Risk_Information_Form.docx")),
            shared_memory.set("phase_1_planning", "risk_analysis_plan_path", os.path.join(phase_output_dir, "Risk_Analysis_Plan.md")),
            shared_memory.set("phase_1_planning", "risk_management_plan_path", os.path.join(phase_output_dir, "Risk_Management_Plan.docx")),
            logging.info(f"Đã lưu các tài liệu Kế hoạch Quản lý Rủi ro và cập nhật shared_memory.")
        )
    )

    # Task: Procurement Tasks
    procurement_tasks = Task(
        description=dedent(f"""
            Tạo 'Procurement_Plan.docx' dựa trên 'WBS.docx' và 'Cost_Estimation_Worksheet.md'.
            Kế hoạch mua sắm cần chi tiết các vật tư, dịch vụ cần mua, lịch trình và quy trình mua sắm.
            --- WBS: {read_file_content(shared_memory.get("phase_1_planning", "wbs_document_path"))}
            --- Cost Estimation Worksheet: {read_file_content(shared_memory.get("phase_1_planning", "cost_estimation_worksheet_path"))}
        """),
        expected_output="Tài liệu tiếng Việt 'Procurement_Plan.docx' chi tiết.",
        agent=planning_orchestrator_agent,
        context=[wbs_task, costing_tasks],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Procurement Tasks ---"),
            process_and_save_docx(str(output), os.path.join(phase_output_dir, "Procurement_Plan.docx"), "Kế hoạch Mua sắm"),
            shared_memory.set("phase_1_planning", "procurement_plan_path", os.path.join(phase_output_dir, "Procurement_Plan.docx")),
            logging.info(f"Đã lưu Procurement_Plan.docx và cập nhật shared_memory.")
        )
    )

    # Task: Configuration Management Tasks
    config_mgmt_tasks = Task(
        description=dedent(f"""
            Tạo 'Configuration_Management_Plan.docx' dựa trên 'Project_Plan.docx', 'WBS.docx' và 'Org_Chart.md'.
            Kế hoạch quản lý cấu hình cần định rõ cách thức quản lý các thay đổi, phiên bản và baseline của dự án.
            --- Project Plan: {read_file_content(shared_memory.get("phase_1_planning", "project_plan_path"))}
            --- WBS: {read_file_content(shared_memory.get("phase_1_planning", "wbs_document_path"))}
            --- Org Chart: {read_file_content(shared_memory.get("phase_1_planning", "org_chart_path"))}
        """),
        expected_output="Tài liệu tiếng Việt 'Configuration_Management_Plan.docx' chi tiết.",
        agent=planning_orchestrator_agent,
        context=[project_plan_task, wbs_task, org_chart_tasks],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Configuration Management Tasks ---"),
            process_and_save_docx(str(output), os.path.join(phase_output_dir, "Configuration_Management_Plan.docx"), "Kế hoạch Quản lý Cấu hình"),
            shared_memory.set("phase_1_planning", "configuration_management_plan_path", os.path.join(phase_output_dir, "Configuration_Management_Plan.docx")),
            logging.info(f"Đã lưu Configuration_Management_Plan.docx và cập nhật shared_memory.")
        )
    )

    # Task: Project Planning Validation
    project_planning_validation_task = Task(
        description=dedent(f"""
            Đánh giá kỹ lưỡng tất cả các tài liệu Kế hoạch Dự án vừa được tạo:
            - 'Project_Plan.docx'
            - 'PMO_Checklist.md'
            - 'COBIT_Checklist.md'
            - 'QA_Checklist.md'
            - 'Statement_of_Work.docx'
            - 'Project_Approval_Document.docx'
            - 'Cost_Estimation_Worksheet.md'
            - 'Development_Estimation.md'
            - 'Capex_Opex_Comparison.md'
            - 'Org_Chart.md'
            - 'Roles_Responsibilities_Matrix.md'
            - 'Approvals_Matrix.md'
            - 'WBS.docx'
            - 'WBS_Dictionary.docx'
            - 'WBS_Resource_Template.docx'
            - 'Project_Plan.txt' (XML simulation)
            - 'WBS_Diagram.png' (existence check if generated)
            - 'Risk_Information_Form.docx'
            - 'Risk_Analysis_Plan.md'
            - 'Risk_Management_Plan.docx'
            - 'Procurement_Plan.docx'
            - 'Configuration_Management_Plan.docx'

            Kiểm tra tính hoàn chỉnh, khả thi, nhất quán với Project Charter và các yêu cầu ban đầu.
            Tạo một báo cáo 'Validation_Report_Phase_1.md' tóm tắt kết quả đánh giá,
            liệt kê các điểm cần cải thiện nếu có và xác nhận việc hoàn thành giai đoạn.
        """),
        expected_output="Tài liệu tiếng Việt 'Validation_Report_Phase_1.md' tóm tắt kết quả đánh giá giai đoạn lập kế hoạch.",
        agent=project_manager_agent,
        context=[
            project_plan_task,
            pmo_tasks,
            sow_tasks,
            approval_tasks,
            costing_tasks,
            org_chart_tasks,
            wbs_task,
            risk_plan_task,
            procurement_tasks,
            config_mgmt_tasks
        ],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Project Planning Validation Task ---"),
            write_output(os.path.join(phase_output_dir, "Validation_Report_Phase_1.md"), str(output)),
            shared_memory.set("phase_1_planning", "validation_report_path", os.path.join(phase_output_dir, "Validation_Report_Phase_1.md")),
            logging.info(f"Đã lưu Validation_Report_Phase_1.md và cập nhật shared_memory.")
        )
    )

    # Task: Research Planning Best Practices
    research_planning_best_practices = Task(
        description=dedent(f"""
            Nghiên cứu các phương pháp hay nhất (best practices) và tiêu chuẩn ngành liên quan đến giai đoạn lập kế hoạch dự án
            để hỗ trợ các agent khác. Ví dụ: cấu trúc điển hình của WBS, ước tính, quản lý rủi ro trong kế hoạch dự án.
            --- Project Plan: {read_file_content(shared_memory.get("phase_1_planning", "project_plan_path"))}
        """),
        expected_output="Tài liệu tiếng Việt 'Planning_Research_Summary.md' tổng hợp các kiến thức nghiên cứu hữu ích.",
        agent=researcher_agent,
        context=[project_plan_task],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Research Planning Best Practices Task ---"),
            write_output(os.path.join(phase_output_dir, "Planning_Research_Summary.md"), str(output)),
            shared_memory.set("phase_1_planning", "research_summary_path", os.path.join(phase_output_dir, "Planning_Research_Summary.md")),
            logging.info(f"Đã lưu Planning_Research_Summary.md và cập nhật shared_memory.")
        )
    )

    return [
        project_plan_task,
        pmo_tasks,
        sow_tasks,
        approval_tasks,
        costing_tasks,
        org_chart_tasks,
        wbs_task,
        risk_plan_task,
        procurement_tasks,
        config_mgmt_tasks,
        project_planning_validation_task,
        research_planning_best_practices
    ]