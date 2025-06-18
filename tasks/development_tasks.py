import os
import logging
import re
from crewai import Task
from textwrap import dedent
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches
import graphviz

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def process_and_create_dev_guidelines_doc(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho Hướng dẫn Phát triển ---")
    doc = Document()
    doc.add_heading('Tài liệu Hướng dẫn Phát triển', level=1)

    dot_blocks = re.findall(r'```dot\s*([\s\S]*?)\s*```', task_output)
    text_content = re.sub(r'```dot\s*[\s\S]*?```', '', task_output).strip()

    guidelines_text = text_content # Giả định phần còn lại là văn bản
    component_diagram_dot_code = ""

    if len(dot_blocks) >= 1:
        component_diagram_dot_code = dot_blocks[0]
        logging.info("Đã trích xuất mã DOT cho Component Diagram.")

    # Thêm nội dung văn bản
    if guidelines_text:
        doc.add_heading('Tiêu chuẩn Mã hóa và Hướng dẫn Phát triển', level=2)
        doc.add_paragraph(guidelines_text)
    else:
        doc.add_paragraph("Không có nội dung hướng dẫn phát triển dạng văn bản được tạo.")

    # Đảm bảo thư mục con tồn tại cho output của phase này
    phase_output_dir = os.path.join(output_base_dir_param, "4_development")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 4: {phase_output_dir}")

    # Render Component Diagram
    if component_diagram_dot_code:
        try:
            graph_component = graphviz.Source(component_diagram_dot_code, format='png', engine='dot')
            component_img_path = os.path.join(phase_output_dir, "component_diagram.png")
            graph_component.render(component_img_path.rsplit('.', 1)[0], view=False, cleanup=True)
            doc.add_heading('Component Diagram', level=2)
            doc.add_picture(component_img_path, width=Inches(6.0))
            logging.info(f"Đã tạo và chèn Component Diagram vào tài liệu: {component_img_path}")
        except Exception as e:
            logging.error(f"Lỗi khi tạo Component Diagram: {e}", exc_info=True)
            doc.add_paragraph(f"Không thể tạo Component Diagram do lỗi: {e}\nMã DOT thất bại:\n```dot\n{component_diagram_dot_code}\n```")
    else:
        doc.add_paragraph("Agent không tạo ra mã DOT cho Component Diagram.")

    final_doc_path = os.path.join(phase_output_dir, "Coding_Standards_and_Guidelines_with_Diagram.docx")
    doc.save(final_doc_path)
    logging.info(f"Đã lưu tài liệu Coding_Standards_and_Guidelines_with_Diagram.docx vào {final_doc_path}")

    shared_memory.set("phase_4_development", "coding_standards_document_path", final_doc_path)
    logging.info(f"--- Hoàn thành xử lý output cho Hướng dẫn Phát triển ---")


def create_development_tasks(lead_software_engineer_agent, project_manager_agent, researcher_agent, output_base_dir):
    """
    Tạo các task liên quan đến giai đoạn Phát triển.
    Args:
        lead_software_engineer_agent: Agent chính cho Development.
        project_manager_agent, researcher_agent: Các agent chung.
        output_base_dir: Đường dẫn thư mục base.
    Returns:
        list: Danh sách các Task đã tạo.
    """
    system_design_doc = shared_memory.get("phase_3_design", "system_architecture") # Assuming this maps to System_Architecture.docx
    db_design_doc = shared_memory.get("phase_3_design", "database_design_document")
    api_design_doc = shared_memory.get("phase_3_design", "api_design_document")
    hld_doc = shared_memory.get("phase_3_design", "hld") # Assuming this maps to High_Level_Design.docx
    lld_doc = shared_memory.get("phase_3_design", "lld") # Assuming this maps to Low_Level_Design.docx
    technical_requirements_doc = shared_memory.get("phase_3_design", "technical_requirements_document") # Assuming this maps to Technical_Requirements_Document.docx
    coding_guidelines_doc_path = shared_memory.get("phase_4_development", "coding_standards_document_path") # Output from dev_standards_tasks

    # Mock inputs if not available from previous phases for demonstration
    # In a real scenario, these would come from actual files or previous task outputs
    configuration_management_plan_doc = "Configuration_Management_Plan.docx (Mock Content)"
    project_plan_xml = "Project_Plan.xml (Mock Content)"
    integration_plan_md = "Integration_Plan.md (Mock Content)"
    coding_guidelines_md = "Coding_Guidelines.md (Mock Content)" # This will be the output from dev_standards_task

    design_context_for_dev = f"System Design Document: {system_design_doc[:500] if system_design_doc else 'N/A'}\n" \
                             f"Database Design (Content Snippet): {db_design_doc[:500] if db_design_doc else 'N/A'}\n" \
                             f"API Design (Content Snippet): {api_design_doc[:500] if api_design_doc else 'N/A'}\n" \
                             f"HLD (Content Snippet): {hld_doc[:500] if hld_doc else 'N/A'}\n" \
                             f"LLD (Content Snippet): {lld_doc[:500] if lld_doc else 'N/A'}\n" \
                             f"Technical Requirements (Content Snippet): {technical_requirements_doc[:500] if technical_requirements_doc else 'N/A'}"
    if not any([system_design_doc, db_design_doc, api_design_doc, hld_doc, lld_doc, technical_requirements_doc]):
        logging.warning("Design documents missing for development tasks.")
        design_context_for_dev = "Không có tài liệu thiết kế nào được tìm thấy."


    # Tạo thư mục con cho Phase 4 Development (nếu chưa có)
    phase_output_dir = os.path.join(output_base_dir, "4_development")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 4: {phase_output_dir}")


    # dev_standards_tasks.py
    # Inputs: Configuration_Management_Plan.docx, Project_Plan.xml
    # Outputs: Development_Standards_Document.md, Coding_Guidelines.md
    # NOTE: The current callback 'process_and_create_dev_guidelines_doc' combines these into a .docx with diagram.
    # The expected_output is adjusted to reflect the actual file created by the callback.
    dev_standards_task = Task(
        description=dedent(f"""
            Dựa trên tài liệu 'Configuration_Management_Plan.docx' và 'Project_Plan.xml', tạo ra một bộ tiêu chuẩn mã hóa (Coding Standards) và hướng dẫn phát triển (Development Guidelines)
            cho dự án. Bao gồm quy tắc đặt tên, cấu trúc thư mục, quy ước comment, và các thực hành tốt nhất cho ngôn ngữ/framework được chọn.
            Bạn cũng phải **tạo mã nguồn Graphviz DOT** cho một Component Diagram đơn giản minh họa các module chính hoặc layer của ứng dụng
            dựa trên thiết kế đã có.
            Cấu trúc đầu ra của bạn phải là văn bản hướng dẫn, sau đó là mã DOT cho Component Diagram
            (trong '```dot\\n...\\n```').
            --- Inputs: Configuration Management Plan: {configuration_management_plan_doc}, Project Plan: {project_plan_xml}
            --- Design Context: {design_context_for_dev}
        """),
        expected_output=dedent("""Một chuỗi văn bản (string) bao gồm:
            1. Phần mô tả Tiêu chuẩn Mã hóa và Hướng dẫn Phát triển.
            2. Tiếp theo là mã Graphviz DOT cho Component Diagram được bọc trong '```dot\\n...\\n```'.
            Đảm bảo mã DOT đúng cú pháp để có thể render thành hình ảnh. (Sẽ được lưu thành 'Coding_Standards_and_Guidelines_with_Diagram.docx')"""),
        agent=lead_software_engineer_agent,
        callback=lambda output: process_and_create_dev_guidelines_doc(str(output), output_base_dir)
    )

    # source_control_tasks.py
    # Inputs: Development_Standards_Document.md, Coding_Guidelines.md (These are outputs from dev_standards_task)
    # Outputs: Version_Control_Plan.md, Source_Code_Repository_Checklist.md
    def source_control_callback(output):
        logging.info(f"--- Hoàn thành Source Control Task ---")
        parts = output.raw.split("--- SOURCE_CODE_REPOSITORY_CHECKLIST ---")
        version_control_content = parts[0].strip()
        repository_checklist_content = parts[1].strip() if len(parts) > 1 else "Không có checklist repository được tạo."

        write_output(os.path.join(phase_output_dir, "Version_Control_Plan.md"), version_control_content)
        shared_memory.set("phase_4_development", "version_control_plan", version_control_content)
        logging.info(f"Đã lưu Version_Control_Plan.md và cập nhật shared_memory.")

        write_output(os.path.join(phase_output_dir, "Source_Code_Repository_Checklist.md"), repository_checklist_content)
        shared_memory.set("phase_4_development", "source_code_repository_checklist", repository_checklist_content)
        logging.info(f"Đã lưu Source_Code_Repository_Checklist.md và cập nhật shared_memory.")

    source_control_task = Task(
        description=dedent(f"""
            Phát triển chiến lược quản lý mã nguồn (Source Control Strategy) cho dự án.
            Xác định công cụ (Git), quy trình branching (ví dụ: Gitflow), và quy tắc commit.
            Tạo thêm một checklist cho việc cấu hình repository mã nguồn.
            --- Inputs: Development Standards & Guidelines (from previous task)
            --- Context: {dev_standards_task.output.raw_output if dev_standards_task.output else 'Chưa có Standards và Guidelines.'}
        """),
        expected_output=dedent("""Một chuỗi văn bản (string) bao gồm:
            1. Nội dung của Version_Control_Plan.md.
            2. Tiếp theo là dấu phân cách '--- SOURCE_CODE_REPOSITORY_CHECKLIST ---'.
            3. Nội dung của Source_Code_Repository_Checklist.md."""),
        agent=lead_software_engineer_agent,
        context=[dev_standards_task],
        callback=source_control_callback
    )

    # integration_tasks.py
    # Inputs: High-Level_Design.docx, API_Design_Document.yaml
    # Outputs: Integration_Plan.md, Unit_Test_Scripts_Template.txt
    def integration_callback(output):
        logging.info(f"--- Hoàn thành Integration Task ---")
        parts = output.raw.split("--- UNIT_TEST_SCRIPTS_TEMPLATE ---")
        integration_plan_content = parts[0].strip()
        unit_test_template_content = parts[1].strip() if len(parts) > 1 else "Không có template test unit được tạo."

        write_output(os.path.join(phase_output_dir, "Integration_Plan.md"), integration_plan_content)
        shared_memory.set("phase_4_development", "integration_plan", integration_plan_content)
        logging.info(f"Đã lưu Integration_Plan.md và cập nhật shared_memory.")

        write_output(os.path.join(phase_output_dir, "Unit_Test_Scripts_Template.txt"), unit_test_template_content)
        shared_memory.set("phase_4_development", "unit_test_scripts_template", unit_test_template_content)
        logging.info(f"Đã lưu Unit_Test_Scripts_Template.txt và cập nhật shared_memory.")

    integration_task = Task(
        description=dedent(f"""
            Dựa trên 'High-Level_Design.docx' và 'API_Design_Document.yaml', tạo kế hoạch tích hợp (Integration Plan) chi tiết.
            Xác định các giai đoạn tích hợp, chiến lược, và trách nhiệm.
            Đồng thời, tạo một template cho các script kiểm thử đơn vị (Unit Test Scripts Template).
            --- Inputs: HLD: {hld_doc[:500] if hld_doc else 'N/A'}, API Design: {api_design_doc[:500] if api_design_doc else 'N/A'}
            --- Context: {source_control_task.output.raw_output if source_control_task.output else 'Chưa có Source Control Strategy.'}
        """),
        expected_output=dedent("""Một chuỗi văn bản (string) bao gồm:
            1. Nội dung của Integration_Plan.md.
            2. Tiếp theo là dấu phân cách '--- UNIT_TEST_SCRIPTS_TEMPLATE ---'.
            3. Nội dung của Unit_Test_Scripts_Template.txt."""),
        agent=lead_software_engineer_agent,
        context=[source_control_task],
        callback=integration_callback
    )

    # build_tasks.py
    # Inputs: System_Architecture.docx, Integration_Plan.md, Configuration_Management_Plan.docx
    # Outputs: Build_and_Deployment_Plan.docx, Build_Verification_Report_Template.docx
    def build_callback(output):
        logging.info(f"--- Hoàn thành Build Task ---")
        parts = output.raw.split("--- BUILD_VERIFICATION_REPORT_TEMPLATE ---")
        build_deploy_plan_content = parts[0].strip()
        bvr_template_content = parts[1].strip() if len(parts) > 1 else "Không có template báo cáo xác minh build được tạo."

        write_output(os.path.join(phase_output_dir, "Build_and_Deployment_Plan.docx"), build_deploy_plan_content)
        shared_memory.set("phase_4_development", "build_deployment_plan", build_deploy_plan_content)
        logging.info(f"Đã lưu Build_and_Deployment_Plan.docx và cập nhật shared_memory.")

        write_output(os.path.join(phase_output_dir, "Build_Verification_Report_Template.docx"), bvr_template_content)
        shared_memory.set("phase_4_development", "build_verification_report_template", bvr_template_content)
        logging.info(f"Đã lưu Build_Verification_Report_Template.docx và cập nhật shared_memory.")

    build_task = Task(
        description=dedent(f"""
            Dựa trên 'System_Architecture.docx', 'Integration_Plan.md', và 'Configuration_Management_Plan.docx',
            xây dựng tài liệu kế hoạch Build và Deployment chi tiết.
            Bao gồm các bước build, môi trường, công cụ, và quy trình triển khai.
            Đồng thời, tạo một template cho báo cáo xác minh build (Build Verification Report Template).
            --- Inputs: System Architecture: {system_design_doc[:500] if system_design_doc else 'N/A'},
                         Integration Plan: {integration_plan_md},
                         Configuration Management Plan: {configuration_management_plan_doc}
            --- Context: {integration_task.output.raw_output if integration_task.output else 'Chưa có Integration Plan.'}
        """),
        expected_output=dedent("""Một chuỗi văn bản (string) bao gồm:
            1. Nội dung của Build_and_Deployment_Plan.docx.
            2. Tiếp theo là dấu phân cách '--- BUILD_VERIFICATION_REPORT_TEMPLATE ---'.
            3. Nội dung của Build_Verification_Report_Template.docx."""),
        agent=lead_software_engineer_agent,
        context=[integration_task],
        callback=build_callback
    )

    # dev_docs_tasks.py
    # Inputs: Low-Level_Design.docx
    # Outputs: Source_Code_Documentation_Template.md, Development_Progress_Report_Template.docx, Middleware_Documentation.md
    def dev_docs_callback(output):
        logging.info(f"--- Hoàn thành Development Documentation Task ---")
        parts = output.raw.split("--- DEVELOPMENT_PROGRESS_REPORT_TEMPLATE ---")
        source_code_doc_template = parts[0].strip()
        remaining_content = parts[1].strip() if len(parts) > 1 else ""
        parts_2 = remaining_content.split("--- MIDDLEWARE_DOCUMENTATION ---")
        dev_progress_report_template = parts_2[0].strip()
        middleware_doc_content = parts_2[1].strip() if len(parts_2) > 1 else "Không có tài liệu middleware được tạo."

        write_output(os.path.join(phase_output_dir, "Source_Code_Documentation_Template.md"), source_code_doc_template)
        shared_memory.set("phase_4_development", "source_code_doc_template", source_code_doc_template)
        logging.info(f"Đã lưu Source_Code_Documentation_Template.md và cập nhật shared_memory.")

        write_output(os.path.join(phase_output_dir, "Development_Progress_Report_Template.docx"), dev_progress_report_template)
        shared_memory.set("phase_4_development", "dev_progress_report_template", dev_progress_report_template)
        logging.info(f"Đã lưu Development_Progress_Report_Template.docx và cập nhật shared_memory.")

        write_output(os.path.join(phase_output_dir, "Middleware_Documentation.md"), middleware_doc_content)
        shared_memory.set("phase_4_development", "middleware_documentation", middleware_doc_content)
        logging.info(f"Đã lưu Middleware_Documentation.md và cập nhật shared_memory.")

    dev_docs_task = Task(
        description=dedent(f"""
            Dựa trên 'Low-Level_Design.docx', tạo các tài liệu hỗ trợ phát triển:
            1. Template tài liệu mã nguồn (Source Code Documentation Template).
            2. Template báo cáo tiến độ phát triển (Development Progress Report Template).
            3. Tài liệu về Middleware được sử dụng (Middleware Documentation).
            --- Inputs: LLD: {lld_doc[:500] if lld_doc else 'N/A'}
            --- Context: {build_task.output.raw_output if build_task.output else 'Chưa có Build Plan.'}
        """),
        expected_output=dedent("""Một chuỗi văn bản (string) bao gồm:
            1. Nội dung của Source_Code_Documentation_Template.md.
            2. Tiếp theo là dấu phân cách '--- DEVELOPMENT_PROGRESS_REPORT_TEMPLATE ---'.
            3. Nội dung của Development_Progress_Report_Template.docx.
            4. Tiếp theo là dấu phân cách '--- MIDDLEWARE_DOCUMENTATION ---'.
            5. Nội dung của Middleware_Documentation.md."""),
        agent=lead_software_engineer_agent,
        context=[build_task],
        callback=dev_docs_callback
    )

    # code_review_tasks.py
    # Inputs: Coding_Guidelines.md, Technical_Requirements_Document.docx
    # Outputs: Code_Review_Checklist.md
    def code_review_callback(output):
        logging.info(f"--- Hoàn thành Code Review Task ---")
        write_output(os.path.join(phase_output_dir, "Code_Review_Checklist.md"), str(output))
        shared_memory.set("phase_4_development", "code_review_checklist", str(output))
        logging.info(f"Đã lưu Code_Review_Checklist.md và cập nhật shared_memory.")

    code_review_task = Task(
        description=dedent(f"""
            Dựa trên 'Coding_Guidelines.md' và 'Technical_Requirements_Document.docx',
            tạo một checklist kiểm tra mã nguồn (Code Review Checklist).
            Checklist này sẽ được sử dụng để đảm bảo chất lượng mã nguồn và tuân thủ tiêu chuẩn.
            --- Inputs: Coding Guidelines: {coding_guidelines_doc_path if coding_guidelines_doc_path else 'N/A'},
                         Technical Requirements: {technical_requirements_doc[:500] if technical_requirements_doc else 'N/A'}
            --- Context: {dev_docs_task.output.raw_output if dev_docs_task.output else 'Chưa có Development Documents.'}
        """),
        expected_output="Tài liệu tiếng Việt 'Code_Review_Checklist.md' mô tả checklist kiểm tra mã nguồn.",
        agent=lead_software_engineer_agent,
        context=[dev_docs_task], # Context should be from dev_standards_task or its output for coding_guidelines
        callback=code_review_callback
    )


    # Task: Development Quality Gate (Project Manager)
    def development_validation_callback(output):
        logging.info(f"--- Hoàn thành Development Validation Task ---")
        write_output(os.path.join(phase_output_dir, "Validation_Report_Phase_4.md"), str(output))
        shared_memory.set("phase_4_development", "validation_report", str(output))
        logging.info(f"Đã lưu Validation_Report_Phase_4.md và cập nhật shared_memory.")

    development_validation_task = Task(
        description=dedent(f"""
            Đánh giá các tài liệu phát triển (Coding Standards, Source Control Strategy, Integration Plan, Build/Deployment Process, Development Docs, Code Review Checklist).
            Kiểm tra tính hoàn chỉnh, khả thi, tuân thủ tiêu chuẩn và yêu cầu thiết kế.
            Tạo báo cáo 'Validation_Report_Phase_4.md' tóm tắt kết quả đánh giá,
            liệt kê các điểm cần cải thiện nếu có và xác nhận hoàn thành giai đoạn.
        """),
        expected_output="Tài liệu tiếng Việt 'Validation_Report_Phase_4.md' tóm tắt kết quả đánh giá giai đoạn phát triển.",
        agent=project_manager_agent,
        context=[dev_standards_task, source_control_task, integration_task, build_task, dev_docs_task, code_review_task],
        callback=development_validation_callback
    )

    # Task: Research Development Best Practices (Researcher)
    def research_development_best_practices_callback(output):
        logging.info(f"--- Hoàn thành Research Development Best Practices Task ---")
        write_output(os.path.join(phase_output_dir, "Development_Research_Summary.md"), str(output))
        shared_memory.set("phase_4_development", "research_summary", str(output))
        logging.info(f"Đã lưu Development_Research_Summary.md và cập nhật shared_memory.")

    research_development_best_practices_task = Task(
        description=dedent(f"""
            Nghiên cứu các phương pháp hay nhất (best practices) trong phát triển phần mềm
            (ví dụ: Clean Code, TDD, DDD, kiến trúc phần mềm, quản lý dependency).
            Tổng hợp kiến thức hỗ trợ các agent khác.
            --- Design Context: {design_context_for_dev}
        """),
        expected_output="Tài liệu tiếng Việt 'Development_Research_Summary.md' tổng hợp các kiến thức nghiên cứu hữu ích.",
        agent=researcher_agent,
        context=[dev_standards_task], # Context could be broader, but starts with standards.
        callback=research_development_best_practices_callback
    )

    return [
        dev_standards_task, # dev_standards_tasks.py
        source_control_task, # source_control_tasks.py
        integration_task, # integration_tasks.py
        build_task, # build_tasks.py
        dev_docs_task, # dev_docs_tasks.py
        code_review_task, # code_review_tasks.py
        development_validation_task,
        research_development_best_practices_task
    ]