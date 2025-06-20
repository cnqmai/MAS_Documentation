import os
import logging
import re
from crewai import Task
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def clean_text_for_docx(text: str) -> str:
    """
    Loại bỏ các ký tự có thể gây lỗi khi ghi vào file .docx hoặc làm phẳng Markdown cơ bản.
    """
    clean_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text).strip()
    clean_text = re.sub(r'#{1,6}\s*', '', clean_text)
    clean_text = clean_text.replace('**', '').replace('__', '')
    clean_text = clean_text.replace('*', '').replace('_', '')
    clean_text = re.sub(r'^- ', '', clean_text, flags=re.MULTILINE)
    clean_text = re.sub(r'^\d+\.\s*', '', clean_text, flags=re.MULTILINE)
    return clean_text

def process_and_save_docx(task_output: str, file_path: str, heading: str):
    doc = Document()
    doc.add_heading(heading, level=1)
    cleaned_output = clean_text_for_docx(str(task_output))
    for paragraph_text in cleaned_output.split('\n'):
        if paragraph_text.strip():
            doc.add_paragraph(paragraph_text)
    doc.save(file_path)
    logging.info(f"Đã lưu {os.path.basename(file_path)}.")

def process_maintenance_plan_documents(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho các tài liệu Kế hoạch Bảo trì ---")
    phase_output_dir = os.path.join(output_base_dir_param, "7_maintenance")
    os.makedirs(phase_output_dir, exist_ok=True)

    plan_content = re.search(r'```maintenance_plan\s*([\s\S]*?)\s*```', task_output)
    checklist_content = re.search(r'```checklist\s*([\s\S]*?)\s*```', task_output)
    sla_policy_content = re.search(r'```sla_policy\s*([\s\S]*?)\s*```', task_output)
    patch_guide_content = re.search(r'```patch_guide\s*([\s\S]*?)\s*```', task_output)

    plan_doc_path = os.path.join(phase_output_dir, "Maintenance_and_Support_Plan.docx")
    plan_text = plan_content.group(1).strip() if plan_content else "Không tìm thấy nội dung Kế hoạch Bảo trì và Hỗ trợ."
    process_and_save_docx(plan_text, plan_doc_path, "Kế hoạch Bảo trì và Hỗ trợ")
    shared_memory.set("phase_7_maintenance", "maintenance_and_support_plan_path", plan_doc_path)
    logging.info(f"Đã lưu Maintenance_and_Support_Plan.docx và cập nhật shared_memory.")

    checklist_md_path = os.path.join(phase_output_dir, "Maintenance_Checklist.md")
    checklist_text = checklist_content.group(1).strip() if checklist_content else "Không tìm thấy nội dung Danh sách kiểm tra bảo trì."
    write_output(checklist_md_path, checklist_text)
    shared_memory.set("phase_7_maintenance", "maintenance_checklist_path", checklist_md_path)
    logging.info(f"Đã lưu Maintenance_Checklist.md và cập nhật shared_memory.")

    sla_policy_doc_path = os.path.join(phase_output_dir, "SLA_and_Warranty_Policies.docx")
    sla_policy_text = sla_policy_content.group(1).strip() if sla_policy_content else "Không tìm thấy nội dung Chính sách SLA và Bảo hành."
    process_and_save_docx(sla_policy_text, sla_policy_doc_path, "Chính sách SLA và Bảo hành")
    shared_memory.set("phase_7_maintenance", "sla_and_warranty_policies_path", sla_policy_doc_path)
    logging.info(f"Đã lưu SLA_and_Warranty_Policies.docx và cập nhật shared_memory.")

    patch_guide_md_path = os.path.join(phase_output_dir, "Patch_Management_Guide.md")
    patch_guide_text = patch_guide_content.group(1).strip() if patch_guide_content else "Không tìm thấy nội dung Hướng dẫn quản lý bản vá."
    write_output(patch_guide_md_path, patch_guide_text)
    shared_memory.set("phase_7_maintenance", "patch_management_guide_path", patch_guide_md_path)
    logging.info(f"Đã lưu Patch_Management_Guide.md và cập nhật shared_memory.")

    logging.info(f"--- Hoàn thành xử lý output cho các tài liệu Kế hoạch Bảo trì ---")

def process_post_project_review_documents(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho các tài liệu Đánh giá Sau Dự án ---")
    phase_output_dir = os.path.join(output_base_dir_param, "7_maintenance")
    os.makedirs(phase_output_dir, exist_ok=True)

    survey_content = re.search(r'```survey_questionnaire\s*([\s\S]*?)\s*```', task_output)
    lessons_learned_content = re.search(r'```lessons_learned\s*([\s\S]*?)\s*```', task_output)
    post_review_content = re.search(r'```post_project_review\s*([\s\S]*?)\s*```', task_output)

    survey_doc_path = os.path.join(phase_output_dir, "Post_Project_Survey_Questionnaire.docx")
    survey_text = survey_content.group(1).strip() if survey_content else "Không tìm thấy nội dung mẫu khảo sát."
    process_and_save_docx(survey_text, survey_doc_path, "Bảng Câu hỏi Khảo sát Sau Dự án")
    shared_memory.set("phase_7_maintenance", "post_project_survey_questionnaire_path", survey_doc_path)
    logging.info(f"Đã lưu Post_Project_Survey_Questionnaire.docx và cập nhật shared_memory.")

    lessons_md_path = os.path.join(phase_output_dir, "Lessons_Learned.md")
    lessons_text = lessons_learned_content.group(1).strip() if lessons_learned_content else "Không tìm thấy nội dung Bài học kinh nghiệm."
    write_output(lessons_md_path, lessons_text)
    shared_memory.set("phase_7_maintenance", "lessons_learned_path", lessons_md_path)
    logging.info(f"Đã lưu Lessons_Learned.md và cập nhật shared_memory.")

    review_doc_path = os.path.join(phase_output_dir, "Post_Project_Review.docx")
    review_text = post_review_content.group(1).strip() if post_review_content else "Không tìm thấy nội dung Đánh giá Sau Dự án."
    process_and_save_docx(review_text, review_doc_path, "Đánh giá Sau Dự án")
    shared_memory.set("phase_7_maintenance", "post_project_review_path", review_doc_path)
    logging.info(f"Đã lưu Post_Project_Review.docx và cập nhật shared_memory.")

    logging.info(f"--- Hoàn thành xử lý output cho các tài liệu Đánh giá Sau Dự án ---")

def process_transition_documents(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho các tài liệu Chuyển giao ---")
    phase_output_dir = os.path.join(output_base_dir_param, "7_maintenance")
    os.makedirs(phase_output_dir, exist_ok=True)

    cr_template_content = re.search(r'```change_request_template\s*([\s\S]*?)\s*```', task_output)
    transition_plan_content = re.search(r'```transition_plan\s*([\s\S]*?)\s*```', task_output)
    retirement_plan_content = re.search(r'```retirement_plan\s*([\s\S]*?)\s*```', task_output)

    cr_template_doc_path = os.path.join(phase_output_dir, "Change_Request_Document_(CCR)_Template.docx")
    cr_template_text = cr_template_content.group(1).strip() if cr_template_content else "Không tìm thấy nội dung Mẫu tài liệu yêu cầu thay đổi."
    process_and_save_docx(cr_template_text, cr_template_doc_path, "Mẫu Tài liệu Yêu cầu Thay đổi (CCR)")
    shared_memory.set("phase_7_maintenance", "change_request_template_path", cr_template_doc_path)
    logging.info(f"Đã lưu Change_Request_Document_(CCR)_Template.docx và cập nhật shared_memory.")

    transition_plan_doc_path = os.path.join(phase_output_dir, "Transition_Out_Plan.docx")
    transition_plan_text = transition_plan_content.group(1).strip() if transition_plan_content else "Không tìm thấy nội dung Kế hoạch chuyển giao."
    process_and_save_docx(transition_plan_text, transition_plan_doc_path, "Kế hoạch Chuyển giao")
    shared_memory.set("phase_7_maintenance", "transition_out_plan_path", transition_plan_doc_path)
    logging.info(f"Đã lưu Transition_Out_Plan.docx và cập nhật shared_memory.")

    retirement_plan_doc_path = os.path.join(phase_output_dir, "Product_Retirement_Plan.docx")
    retirement_plan_text = retirement_plan_content.group(1).strip() if retirement_plan_content else "Không tìm thấy nội dung Kế hoạch ngừng sản phẩm."
    process_and_save_docx(retirement_plan_text, retirement_plan_doc_path, "Kế hoạch Ngừng Sản phẩm")
    shared_memory.set("phase_7_maintenance", "product_retirement_plan_path", retirement_plan_doc_path)
    logging.info(f"Đã lưu Product_Retirement_Plan.docx và cập nhật shared_memory.")

    logging.info(f"--- Hoàn thành xử lý output cho các tài liệu Chuyển giao ---")

def process_knowledge_transfer_documents(task_output: str, output_base_dir_param: str):
    logging.info(f"--- Bắt đầu xử lý output cho các tài liệu Chuyển giao Kiến thức ---")
    phase_output_dir = os.path.join(output_base_dir_param, "7_maintenance")
    os.makedirs(phase_output_dir, exist_ok=True)

    dev_knowledge_content = re.search(r'```dev_knowledge\s*([\s\S]*?)\s*```', task_output)
    support_summary_content = re.search(r'```support_summary\s*([\s\S]*?)\s*```', task_output)

    dev_knowledge_md_path = os.path.join(phase_output_dir, "Developer_Knowledge_Transfer_Report.md")
    dev_knowledge_text = dev_knowledge_content.group(1).strip() if dev_knowledge_content else "Không tìm thấy nội dung Báo cáo chuyển giao kiến thức cho nhà phát triển."
    write_output(dev_knowledge_md_path, dev_knowledge_text)
    shared_memory.set("phase_7_maintenance", "developer_knowledge_transfer_report_path", dev_knowledge_md_path)
    logging.info(f"Đã lưu Developer_Knowledge_Transfer_Report.md và cập nhật shared_memory.")

    support_summary_md_path = os.path.join(phase_output_dir, "Global_Application_Support_Summary.md")
    support_summary_text = support_summary_content.group(1).strip() if support_summary_content else "Không tìm thấy nội dung Tóm tắt hỗ trợ ứng dụng toàn cầu."
    write_output(support_summary_md_path, support_summary_text)
    shared_memory.set("phase_7_maintenance", "global_application_support_summary_path", support_summary_md_path)
    logging.info(f"Đã lưu Global_Application_Support_Summary.md và cập nhật shared_memory.")

    logging.info(f"--- Hoàn thành xử lý output cho các tài liệu Chuyển giao Kiến thức ---")

def create_maintenance_tasks(site_reliability_engineer_agent, project_manager_agent, researcher_agent, output_base_dir):
    deployment_plan_doc_path = shared_memory.get("phase_6_deployment", "production_implementation_plan_docx_path") or "N/A"
    monitoring_guide_path = shared_memory.get("phase_6_deployment", "monitoring_and_alerting_setup_guide_path") or "N/A"
    source_code_doc_path = shared_memory.get("phase_4_development", "source_code_document_path") or "N/A"
    api_design_doc_path = shared_memory.get("phase_3_design", "api_design_document_path") or "N/A"

    maintenance_context_base = f"Deployment Plan: {deployment_plan_doc_path}\n" \
                               f"Monitoring Setup Guide: {monitoring_guide_path}"
    if deployment_plan_doc_path == "N/A" and monitoring_guide_path == "N/A":
        logging.warning("Previous phase deployment/monitoring documents missing for maintenance tasks.")
        maintenance_context_base = "Không có tài liệu triển khai/giám sát liên quan nào được tìm thấy."

    knowledge_transfer_context = f"Source Code Documentation: {source_code_doc_path}\n" \
                                 f"API Design Document: {api_design_doc_path}"
    if source_code_doc_path == "N/A" and api_design_doc_path == "N/A":
        logging.warning("Previous phase development/design documents missing for knowledge transfer tasks.")
        knowledge_transfer_context = "Không có tài liệu mã nguồn/thiết kế API liên quan nào được tìm thấy."

    phase_output_dir = os.path.join(output_base_dir, "7_maintenance")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đã đảm bảo thư mục đầu ra cho Phase 7: {phase_output_dir}")

    maintenance_plan_tasks = Task(
        description=(
            f"Tạo các tài liệu sau trong một phản hồi duy nhất, sử dụng các khối mã riêng biệt:\n"
            f"1. **Maintenance_and_Support_Plan.docx**: Kế hoạch chi tiết về bảo trì và hỗ trợ, bao gồm các hoạt động định kỳ, nâng cấp, vá lỗi, quản lý phiên bản và hỗ trợ kỹ thuật.\n"
            f"2. **Maintenance_Checklist.md**: Danh sách kiểm tra các hoạt động bảo trì cần thực hiện.\n"
            f"3. **SLA_and_Warranty_Policies.docx**: Các chính sách về Thỏa thuận mức độ dịch vụ (SLA) và bảo hành cho hệ thống.\n"
            f"4. **Patch_Management_Guide.md**: Hướng dẫn chi tiết quy trình quản lý và triển khai các bản vá lỗi.\n"
            f"Cấu trúc đầu ra của bạn phải như sau:\n"
            f"```maintenance_plan\n[Nội dung Kế hoạch Bảo trì và Hỗ trợ]\n```\n"
            f"```checklist\n[Nội dung Danh sách Kiểm tra Bảo trì]\n```\n"
            f"```sla_policy\n[Nội dung Chính sách SLA và Bảo hành]\n```\n"
            f"```patch_guide\n[Nội dung Hướng dẫn Quản lý Bản vá]\n```\n"
            f"--- Context: {maintenance_context_base}"
        ),
        expected_output=(
            "Một chuỗi văn bản (string) chứa nội dung cho bốn tài liệu được phân tách bằng các khối mã:\n"
            "'- `maintenance_plan` cho Maintenance_and_Support_Plan.docx'\n"
            "'- `checklist` cho Maintenance_Checklist.md'\n"
            "'- `sla_policy` cho SLA_and_Warranty_Policies.docx'\n"
            "'- `patch_guide` cho Patch_Management_Guide.md'"
        ),
        agent=site_reliability_engineer_agent,
        callback=lambda output: process_maintenance_plan_documents(str(output), output_base_dir)
    )

    feedback_review_tasks = Task(
        description=(
            f"Tạo ba tài liệu sau trong một phản hồi duy nhất, sử dụng các khối mã riêng biệt:\n"
            f"1. **Post_Project_Survey_Questionnaire.docx**: Mẫu câu hỏi khảo sát sau dự án để thu thập phản hồi từ các bên liên quan.\n"
            f"2. **Lessons_Learned.md**: Báo cáo tổng hợp các bài học kinh nghiệm rút ra từ toàn bộ dự án.\n"
            f"3. **Post_Project_Review.docx**: Tài liệu đánh giá tổng thể dự án sau khi hoàn thành, so sánh với mục tiêu ban đầu và đề xuất cải tiến.\n"
            f"Cấu trúc đầu ra của bạn phải như sau:\n"
            f"```survey_questionnaire\n[Nội dung Mẫu khảo sát]\n```\n"
            f"```lessons_learned\n[Nội dung Bài học kinh nghiệm]\n```\n"
            f"```post_project_review\n[Nội dung Đánh giá Sau Dự án]\n```\n"
            f"--- Context: Toàn bộ thông tin từ shared_memory về các giai đoạn trước để đánh giá toàn diện."
        ),
        expected_output=(
            "Một chuỗi văn bản (string) chứa nội dung cho ba tài liệu được phân tách bằng các khối mã:\n"
            "'- `survey_questionnaire` cho Post_Project_Survey_Questionnaire.docx'\n"
            "'- `lessons_learned` cho Lessons_Learned.md'\n"
            "'- `post_project_review` cho Post_Project_Review.docx'"
        ),
        agent=project_manager_agent,
        context=[maintenance_plan_tasks],
        callback=lambda output: process_post_project_review_documents(str(output), output_base_dir)
    )

    transition_tasks = Task(
        description=(
            f"Tạo ba tài liệu sau trong một phản hồi duy nhất, sử dụng các khối mã riêng biệt:\n"
            f"1. **Change_Request_Document_(CCR)_Template.docx**: Mẫu tài liệu yêu cầu thay đổi (Change Control Request - CCR) cho việc quản lý thay đổi trong tương lai.\n"
            f"2. **Transition_Out_Plan.docx**: Kế hoạch chuyển giao hệ thống hoặc trách nhiệm hỗ trợ cho một bên khác nếu có.\n"
            f"3. **Product_Retirement_Plan.docx**: Kế hoạch chi tiết cho việc ngừng hoặc loại bỏ sản phẩm/hệ thống khi không còn cần thiết.\n"
            f"Cấu trúc đầu ra của bạn phải như sau:\n"
            f"```change_request_template\n[Nội dung Mẫu Yêu cầu Thay đổi]\n```\n"
            f"```transition_plan\n[Nội dung Kế hoạch Chuyển giao]\n```\n"
            f"```retirement_plan\n[Nội dung Kế hoạch Ngừng Sản phẩm]\n```\n"
            f"--- Context: {maintenance_context_base}"
        ),
        expected_output=(
            "Một chuỗi văn bản (string) chứa nội dung cho ba tài liệu được phân tách bằng các khối mã:\n"
            "'- `change_request_template` cho Change_Request_Document_(CCR)_Template.docx'\n"
            "'- `transition_plan` cho Transition_Out_Plan.docx'\n"
            "'- `retirement_plan` cho Product_Retirement_Plan.docx'"
        ),
        agent=site_reliability_engineer_agent,
        context=[maintenance_plan_tasks, feedback_review_tasks],
        callback=lambda output: process_transition_documents(str(output), output_base_dir)
    )

    support_knowledge_tasks = Task(
        description=(
            f"Tạo hai tài liệu sau trong một phản hồi duy nhất, sử dụng các khối mã riêng biệt:\n"
            f"1. **Developer_Knowledge_Transfer_Report.md**: Báo cáo tổng hợp kiến thức kỹ thuật quan trọng cần chuyển giao cho các nhà phát triển hoặc nhóm hỗ trợ cấp cao.\n"
            f"2. **Global_Application_Support_Summary.md**: Tóm tắt các thông tin quan trọng về hỗ trợ ứng dụng trên toàn cầu, bao gồm liên hệ, quy trình và công cụ.\n"
            f"Cấu trúc đầu ra của bạn phải như sau:\n"
            f"```dev_knowledge\n[Nội dung Báo cáo Chuyển giao Kiến thức cho Nhà phát triển]\n```\n"
            f"```support_summary\n[Nội dung Tóm tắt Hỗ trợ Ứng dụng Toàn cầu]\n```\n"
            f"--- Context: {knowledge_transfer_context}"
        ),
        expected_output=(
            "Một chuỗi văn bản (string) chứa nội dung cho hai tài liệu được phân tách bằng các khối mã:\n"
            "'- `dev_knowledge` cho Developer_Knowledge_Transfer_Report.md'\n"
            "'- `support_summary` cho Global_Application_Support_Summary.md'"
        ),
        agent=site_reliability_engineer_agent,
        context=[maintenance_plan_tasks],
        callback=lambda output: process_knowledge_transfer_documents(str(output), output_base_dir)
    )

    research_maintenance_best_practices = Task(
        description=(
            f"Nghiên cứu các phương pháp hay nhất (best practices) trong bảo trì và vận hành hệ thống "
            f"(ví dụ: ITIL, Site Reliability Engineering (SRE) practices, incident management). "
            f"Tổng hợp kiến thức hỗ trợ các agent khác.\n"
            f"--- Context: Kế hoạch Bảo trì và Hỗ trợ: {shared_memory.get('phase_7_maintenance', 'maintenance_and_support_plan_path') or 'N/A'}"
        ),
        expected_output="Tài liệu tiếng Việt 'Maintenance_Research_Summary.md' tổng hợp các kiến thức nghiên cứu hữu ích.",
        agent=researcher_agent,
        context=[maintenance_plan_tasks],
        callback=lambda output: (
            logging.info(f"--- Hoàn thành Research Maintenance Best Practices Task ---"),
            write_output(os.path.join(phase_output_dir, "Maintenance_Research_Summary.md"), str(output)),
            shared_memory.set("phase_7_maintenance", "research_summary_path", str(output)),
            logging.info(f"Đã lưu Maintenance_Research_Summary.md và cập nhật shared_memory.")
        )
    )

    return [
        maintenance_plan_tasks,
        feedback_review_tasks,
        transition_tasks,
        support_knowledge_tasks,
        research_maintenance_best_practices
    ]