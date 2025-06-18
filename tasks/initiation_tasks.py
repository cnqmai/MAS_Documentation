import os
import logging
import re
from crewai import Task
from utils.file_writer import write_output
from memory.shared_memory import shared_memory
from docx import Document
from docx.shared import Inches
from textwrap import dedent
import pandas as pd

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def clean_text_for_docx(text: str) -> str:
    """
    Loại bỏ các ký tự có thể gây lỗi khi ghi vào file .docx hoặc làm phẳng Markdown cơ bản.
    """
    clean_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text).strip()
    clean_text = re.sub(r'#{1,6}\s*', '', clean_text)
    clean_text = clean_text.replace('**', '').replace('__', '')
    clean_text = clean_text.replace('*', '').replace('_', '')
    clean_text = re.sub(r'^- ', '', clean_text, flags=re.MULTILINE)
    clean_text = re.sub(r'^\d+\.\s*', '', clean_text, flags=re.MULTILINE)
    return clean_text

def process_and_save_docx(task_output: str, file_path: str, heading: str):
    doc = Document()
    doc.add_heading(heading, level=1)
    cleaned_output = clean_text_for_docx(str(task_output))
    for paragraph_text in cleaned_output.split('\n'):
        if paragraph_text.strip():
            doc.add_paragraph(paragraph_text)
    doc.save(file_path)
    logging.info(f"Đã lưu {os.path.basename(file_path)}.")

def process_and_save_xlsx(task_output: str, file_path: str, sheet_name: str = 'Sheet1'):
    try:
        lines = task_output.strip().split('\n')
        if not lines:
            raise ValueError("Output is empty")
        if ',' in lines[0] or '\t' in lines[0]:
            data = [line.split(',') if ',' in lines[0] else line.split('\t') for line in lines]
            df = pd.DataFrame(data[1:], columns=data[0])
        else:
            if any(':' in line for line in lines):
                data = {}
                for line in lines:
                    if ':' in line:
                        key, value = line.split(':', 1)
                        data[key.strip()] = value.strip()
                df = pd.DataFrame([data]) if data else pd.DataFrame({'Content': lines})
            else:
                df = pd.DataFrame({'Content': lines})
        df.to_excel(file_path, index=False, sheet_name=sheet_name)
        logging.info(f"Đã lưu {os.path.basename(file_path)}.")
    except Exception as e:
        logging.error(f"Lỗi khi lưu file Excel {file_path}: {e}. Lưu dưới dạng Markdown thay thế.")
        write_output(file_path.replace('.xlsx', '.md'), f"Lỗi tạo Excel: {e}\n\nNội dung gốc:\n{task_output}")
    shared_memory.set("phase_0_initiation", os.path.basename(file_path).replace('.', '_').lower() + '_path', file_path)

def create_initiation_tasks(project_manager_agent, researcher_agent, initiation_agent, output_base_dir):
    phase_output_dir = os.path.join(output_base_dir, "0_initiation")
    os.makedirs(phase_output_dir, exist_ok=True)
    logging.info(f"Đảm bảo thư mục output cho Phase 0 tồn tại: {phase_output_dir}")

    def initiation_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'initiation_tasks' ---")
        project_charter_content_match = re.search(r'```project_charter\s*([\s\S]*?)\s*```', str(output))
        business_case_content_match = re.search(r'```business_case\s*([\s\S]*?)\s*```', str(output))
        feasibility_report_content_match = re.search(r'```feasibility_report\s*([\s\S]*?)\s*```', str(output))
        if project_charter_content_match:
            process_and_save_docx(project_charter_content_match.group(1), os.path.join(phase_output_dir, "Project_Charter.docx"), "Điều lệ Dự án")
            shared_memory.set("phase_0_initiation", "project_charter_path", os.path.join(phase_output_dir, "Project_Charter.docx"))
        else:
            logging.warning("Không tìm thấy nội dung Project Charter. File Project_Charter.docx sẽ không được tạo.")
        if business_case_content_match:
            process_and_save_docx(business_case_content_match.group(1), os.path.join(phase_output_dir, "Business_Case.docx"), "Trường hợp Kinh doanh")
            shared_memory.set("phase_0_initiation", "business_case_path", os.path.join(phase_output_dir, "Business_Case.docx"))
        else:
            logging.warning("Không tìm thấy nội dung Business Case. File Business_Case.docx sẽ không được tạo.")
        if feasibility_report_content_match:
            process_and_save_docx(feasibility_report_content_match.group(1), os.path.join(phase_output_dir, "Feasibility_Report.docx"), "Báo cáo Khả thi")
            shared_memory.set("phase_0_initiation", "feasibility_report_path", os.path.join(phase_output_dir, "Feasibility_Report.docx"))
        else:
            logging.warning("Không tìm thấy nội dung Feasibility Report. File Feasibility_Report.docx sẽ không được tạo.")

    initiation_tasks = Task(
        description=dedent(f"""
            Là {getattr(initiation_agent, 'role', 'Chuyên gia Khởi tạo')}, bạn phải xác định và định nghĩa dự án một cách rõ ràng.
            Tạo ba tài liệu quan trọng:
            1.  **Điều lệ Dự án (Project Charter):** Ủy quyền cho Project Manager, phác thảo các mục tiêu cấp cao, phạm vi ban đầu và các bên liên quan chính.
            2.  **Trường hợp Kinh doanh (Business Case):** Mô tả vấn đề kinh doanh, giải pháp được đề xuất, lợi ích mong đợi và phân tích sơ bộ về chi phí-lợi ích.
            3.  **Báo cáo Khả thi (Feasibility Report):** Đánh giá tính khả thi kỹ thuật, vận hành, pháp lý và kinh tế của dự án.
            Sử dụng thông tin thu thập được từ yêu cầu người dùng ban đầu.
            Output của bạn phải chứa 3 phần riêng biệt, mỗi phần được bao bọc bởi các tag độc đáo:
            ```project_charter
            [Nội dung Điều lệ Dự án ở đây]
            ```
            ```business_case
            [Nội dung Trường hợp Kinh doanh ở đây]
            ```
            ```feasibility_report
            [Nội dung Báo cáo Khả thi ở đây]
            ```
            Đảm bảo nội dung đầy đủ và chuyên nghiệp cho từng tài liệu.
            --- Context:
            Yêu cầu Người dùng Ban đầu: {shared_memory.get('phase_0_initiation', 'system_request_summary') or 'Chưa có yêu cầu người dùng ban đầu.'}
        """),
        expected_output="3 tài liệu tiếng Việt: 'Project_Charter.docx', 'Business_Case.docx', 'Feasibility_Report.docx' với nội dung chuyên nghiệp, được trích xuất từ 3 khối đầu ra riêng biệt của bạn.",
        agent=initiation_agent,
        context=[],
        callback=initiation_tasks_callback
    )

    def keyword_generation_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'keyword_generation_tasks' ---")
        filepath = os.path.join(phase_output_dir, "Query_Keywords.txt")
        write_output(filepath, str(output))
        shared_memory.set("phase_0_initiation", "query_keywords_path", filepath)
        logging.info(f"Đã lưu Query_Keywords.txt và cập nhật shared_memory.")

    keyword_generation_tasks = Task(
        description=dedent(f"""
            Dựa trên Điều lệ Dự án và yêu cầu ban đầu, tạo ra các từ khóa và câu truy vấn tìm kiếm cần thiết
            để Researcher Agent có thể thực hiện nghiên cứu sâu hơn về các công nghệ, giải pháp, tiêu chuẩn ngành,
            và các vấn đề tiềm năng liên quan đến dự án.
            Output là một danh sách các từ khóa/câu truy vấn, mỗi mục trên một dòng.
            --- Context:
            Yêu cầu Người dùng Ban đầu: {shared_memory.get('phase_0_initiation', 'system_request_summary') or 'N/A'}
            Điều lệ Dự án: {shared_memory.get('phase_0_initiation', 'project_charter_path') or 'N/A'}
        """),
        expected_output="Tài liệu tiếng Việt 'Query_Keywords.txt' liệt kê các từ khóa và truy vấn tìm kiếm, mỗi từ/câu trên một dòng riêng biệt.",
        agent=researcher_agent,
        context=[initiation_tasks],
        callback=keyword_generation_tasks_callback
    )

    def phase_req_analysis_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'phase_req_analysis_tasks' ---")
        filepath = os.path.join(phase_output_dir, "Phase_Specific_Requirements.docx")
        process_and_save_docx(str(output), filepath, "Yêu cầu Cụ thể Giai đoạn Khởi tạo")
        shared_memory.set("phase_0_initiation", "phase_specific_requirements_path", filepath)
        logging.info(f"Đã lưu Phase_Specific_Requirements.docx và cập nhật shared_memory.")

    phase_req_analysis_tasks = Task(
        description=dedent(f"""
            Là Project Manager, bạn cần phân tích các yêu cầu đã thu thập được từ yêu cầu ban đầu và Điều lệ Dự án.
            Xác định và tài liệu hóa các yêu cầu cụ thể cần được đáp ứng **trong Giai đoạn Khởi tạo (Phase 0) này**.
            Điều này bao gồm các yêu cầu để hoàn thành các deliverables của Phase 0, không phải toàn bộ yêu cầu dự án.
            Output là một tài liệu Word với các yêu cầu được trình bày rõ ràng, có cấu trúc.
            --- Context:
            Yêu cầu Người dùng Ban đầu: {shared_memory.get('phase_0_initiation', 'system_request_summary') or 'N/A'}
            Điều lệ Dự án: {shared_memory.get('phase_0_initiation', 'project_charter_path') or 'N/A'}
        """),
        expected_output="Tài liệu tiếng Việt 'Phase_Specific_Requirements.docx' mô tả chi tiết các yêu cầu cụ thể cho giai đoạn khởi tạo của dự án.",
        agent=project_manager_agent,
        context=[initiation_tasks],
        callback=phase_req_analysis_tasks_callback
    )

    def research_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'research_tasks' ---")
        best_practices_match = re.search(r'```best_practices_report\s*([\s\S]*?)\s*```', str(output))
        ieee_uml_match = re.search(r'```ieee_uml_guidelines\s*([\s\S]*?)\s*```', str(output))
        architecture_references_match = re.search(r'```architecture_references\s*([\s\S]*?)\s*```', str(output))
        template_library_match = re.search(r'```template_library_update\s*([\s\S]*?)\s*```', str(output))
        if best_practices_match:
            filepath = os.path.join(phase_output_dir, "Best_Practices_Report.md")
            write_output(filepath, best_practices_match.group(1))
            shared_memory.set("phase_0_initiation", "best_practices_report_path", filepath)
            logging.info(f"Đã lưu Best_Practices_Report.md.")
        else:
            logging.warning("Không tìm thấy nội dung Best Practices Report.")
        if ieee_uml_match:
            filepath = os.path.join(phase_output_dir, "IEEE_UML_Guidelines.md")
            write_output(filepath, ieee_uml_match.group(1))
            shared_memory.set("phase_0_initiation", "ieee_uml_guidelines_path", filepath)
            logging.info(f"Đã lưu IEEE_UML_Guidelines.md.")
        else:
            logging.warning("Không tìm thấy nội dung IEEE UML Guidelines.")
        if architecture_references_match:
            filepath = os.path.join(phase_output_dir, "Architecture_References.md")
            write_output(filepath, architecture_references_match.group(1))
            shared_memory.set("phase_0_initiation", "architecture_references_path", filepath)
            logging.info(f"Đã lưu Architecture_References.md.")
        else:
            logging.warning("Không tìm thấy nội dung Architecture References.")
        if template_library_match:
            filepath = os.path.join(phase_output_dir, "Template_Library_Update.md")
            write_output(filepath, template_library_match.group(1))
            shared_memory.set("phase_0_initiation", "template_library_update_path", filepath)
            logging.info(f"Đã lưu Template_Library_Update.md.")
        else:
            logging.warning("Không tìm thấy nội dung Template Library Update.")

    research_tasks = Task(
        description=dedent(f"""
            Là {getattr(researcher_agent, 'role', 'Nhà Nghiên cứu')}, bạn hãy thực hiện nghiên cứu chuyên sâu về các thực hành tốt nhất
            và tiêu chuẩn ngành (ví dụ: IEEE cho UML, các nguyên tắc kiến trúc phần mềm) liên quan đến dự án và giai đoạn khởi tạo.
            Đồng thời, xác định và tóm tắt các tài liệu tham khảo kiến trúc quan trọng và đề xuất các cập nhật
            cho thư viện mẫu nếu có.
            Output của bạn phải chứa 4 phần riêng biệt, mỗi phần được bao bọc bởi các tag độc đáo:
            ```best_practices_report
            [Nội dung báo cáo thực hành tốt nhất]
            ```
            ```ieee_uml_guidelines
            [Mô tả các hướng dẫn IEEE/UML chính]
            ```
            ```architecture_references
            [Mô tả các tài liệu/nguồn tham khảo kiến trúc quan trọng]
            ```
            ```template_library_update
            [Đề xuất cập nhật thư viện mẫu]
            ```
            --- Context:
            Từ khóa Truy vấn: {shared_memory.get('phase_0_initiation', 'query_keywords_path') or 'N/A'}
            Yêu cầu Cụ thể Giai đoạn Khởi tạo: {shared_memory.get('phase_0_initiation', 'phase_specific_requirements_path') or 'N/A'}
        """),
        expected_output="4 tài liệu tiếng Việt: 'Best_Practices_Report.md', 'IEEE_UML_Guidelines.md', 'Architecture_References.md', 'Template_Library_Update.md'.",
        agent=researcher_agent,
        context=[keyword_generation_tasks, phase_req_analysis_tasks],
        callback=research_tasks_callback
    )

    def stakeholder_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'stakeholder_tasks' ---")
        stakeholder_list_match = re.search(r'```stakeholder_list\s*([\s\S]*?)\s*```', str(output))
        stakeholder_analysis_match = re.search(r'```stakeholder_analysis\s*([\s\S]*?)\s*```', str(output))
        project_submission_form_match = re.search(r'```project_submission_form\s*([\s\S]*?)\s*```', str(output))
        if stakeholder_list_match:
            process_and_save_xlsx(stakeholder_list_match.group(1), os.path.join(phase_output_dir, "Stakeholder_List.xlsx"), "Danh sách Bên liên quan")
            shared_memory.set("phase_0_initiation", "stakeholder_list_path", os.path.join(phase_output_dir, "Stakeholder_List.xlsx"))
        else:
            logging.warning("Không tìm thấy nội dung Stakeholder List.")
        if stakeholder_analysis_match:
            filepath = os.path.join(phase_output_dir, "Stakeholder_Analysis.md")
            write_output(filepath, stakeholder_analysis_match.group(1))
            shared_memory.set("phase_0_initiation", "stakeholder_analysis_path", filepath)
            logging.info(f"Đã lưu Stakeholder_Analysis.md.")
        else:
            logging.warning("Không tìm thấy nội dung Stakeholder Analysis.")
        if project_submission_form_match:
            process_and_save_docx(project_submission_form_match.group(1), os.path.join(phase_output_dir, "Project_Submission_Form.docx"), "Biểu mẫu Nộp Dự án")
            shared_memory.set("phase_0_initiation", "project_submission_form_path", os.path.join(phase_output_dir, "Project_Submission_Form.docx"))
        else:
            logging.warning("Không tìm thấy nội dung Project Submission Form.")

    stakeholder_tasks = Task(
        description=dedent(f"""
            Là {getattr(initiation_agent, 'role', 'Chuyên gia Khởi tạo')}, bạn hãy xác định và phân tích các bên liên quan chính của dự án.
            Tạo một danh sách các bên liên quan (Stakeholder List) ở định dạng bảng (phù hợp cho Excel),
            thực hiện phân tích các bên liên quan (Stakeholder Analysis) bao gồm mối quan tâm, mức độ ảnh hưởng của họ.
            Đồng thời, tạo một biểu mẫu nộp dự án (Project Submission Form) cơ bản.
            Output của bạn phải chứa 3 phần riêng biệt, mỗi phần được bao bọc bởi các tag độc đáo:
            ```stakeholder_list
            [Nội dung danh sách bên liên quan (định dạng bảng)]
            ```
            ```stakeholder_analysis
            [Nội dung phân tích bên liên quan]
            ```
            ```project_submission_form
            [Nội dung biểu mẫu nộp dự án]
            ```
            --- Context:
            Điều lệ Dự án: {shared_memory.get('phase_0_initiation', 'project_charter_path') or 'N/A'}
        """),
        expected_output="3 tài liệu tiếng Việt: 'Stakeholder_List.xlsx', 'Stakeholder_Analysis.md', 'Project_Submission_Form.docx'.",
        agent=initiation_agent,
        context=[initiation_tasks],
        callback=stakeholder_tasks_callback
    )

    def resourcing_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'resourcing_tasks' ---")
        team_definition_match = re.search(r'```project_team_definition\s*([\s\S]*?)\s*```', str(output))
        resource_plan_match = re.search(r'```project_resource_plan\s*([\s\S]*?)\s*```', str(output))
        if team_definition_match:
            process_and_save_docx(team_definition_match.group(1), os.path.join(phase_output_dir, "Project_Team_Definition.docx"), "Định nghĩa Đội ngũ Dự án")
            shared_memory.set("phase_0_initiation", "project_team_definition_path", os.path.join(phase_output_dir, "Project_Team_Definition.docx"))
        else:
            logging.warning("Không tìm thấy nội dung Project Team Definition.")
        if resource_plan_match:
            process_and_save_xlsx(resource_plan_match.group(1), os.path.join(phase_output_dir, "Project_Resource_Plan.xlsx"), "Kế hoạch Nguồn lực")
            shared_memory.set("phase_0_initiation", "project_resource_plan_path", os.path.join(phase_output_dir, "Project_Resource_Plan.xlsx"))
        else:
            logging.warning("Không tìm thấy nội dung Project Resource Plan.")

    resourcing_tasks = Task(
        description=dedent(f"""
            Là {getattr(initiation_agent, 'role', 'Chuyên gia Khởi tạo')}, bạn hãy lập kế hoạch ban đầu về nguồn lực và đội ngũ cần thiết cho dự án.
            Tạo một tài liệu định nghĩa đội ngũ dự án (Project Team Definition) phác thảo các vai trò, trách nhiệm chính.
            Tạo một kế hoạch nguồn lực dự án (Project Resource Plan) ở định dạng bảng (phù hợp cho Excel)
            liệt kê các loại nguồn lực cần thiết (nhân sự, công cụ, v.v.) và ước tính sơ bộ.
            Output của bạn phải chứa 2 phần riêng biệt, mỗi phần được bao bọc bởi các tag độc đáo:
            ```project_team_definition
            [Nội dung định nghĩa đội ngũ]
            ```
            ```project_resource_plan
            [Nội dung kế hoạch nguồn lực (định dạng bảng)]
            ```
            --- Context:
            Danh sách Bên liên quan: {shared_memory.get('phase_0_initiation', 'stakeholder_list_path') or 'N/A'}
        """),
        expected_output="2 tài liệu tiếng Việt: 'Project_Team_Definition.docx', 'Project_Resource_Plan.xlsx'.",
        agent=initiation_agent,
        context=[stakeholder_tasks],
        callback=resourcing_tasks_callback
    )

    def conops_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'conops_tasks' ---")
        filepath = os.path.join(phase_output_dir, "Concept_of_Operations.docx")
        process_and_save_docx(str(output), filepath, "Khái niệm Vận hành (ConOps)")
        shared_memory.set("phase_0_initiation", "concept_of_operations_path", filepath)
        logging.info(f"Đã lưu Concept_of_Operations.docx và cập nhật shared_memory.")

    conops_tasks = Task(
        description=dedent(f"""
            Là {getattr(initiation_agent, 'role', 'Chuyên gia Khởi tạo')}, bạn hãy phát triển một tài liệu Khái niệm Vận hành (Concept of Operations - ConOps) sơ bộ.
            Tài liệu này nên mô tả cách hệ thống sẽ được sử dụng từ góc độ người dùng,
            các kịch bản hoạt động chính, và môi trường hoạt động dự kiến.
            Output là một tài liệu Word chuyên nghiệp và rõ ràng.
            --- Context:
            Trường hợp Kinh doanh: {shared_memory.get('phase_0_initiation', 'business_case_path') or 'N/A'}
        """),
        expected_output="Tài liệu tiếng Việt 'Concept_of_Operations.docx' mô tả khái niệm vận hành của hệ thống.",
        agent=initiation_agent,
        context=[initiation_tasks],
        callback=conops_tasks_callback
    )

    def risk_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'risk_tasks' ---")
        risk_assessment_match = re.search(r'```risk_assessment\s*([\s\S]*?)\s*```', str(output))
        checklist_match = re.search(r'```initiate_project_checklist\s*([\s\S]*?)\s*```', str(output))
        if risk_assessment_match:
            filepath = os.path.join(phase_output_dir, "Risk_Assessment_Document.md")
            write_output(filepath, risk_assessment_match.group(1))
            shared_memory.set("phase_0_initiation", "risk_assessment_document_path", filepath)
            logging.info(f"Đã lưu Risk_Assessment_Document.md.")
        else:
            logging.warning("Không tìm thấy nội dung Risk Assessment.")
        if checklist_match:
            process_and_save_xlsx(checklist_match.group(1), os.path.join(phase_output_dir, "Initiate_Project_Checklist.xlsx"), "Danh sách Kiểm tra Khởi tạo")
            shared_memory.set("phase_0_initiation", "initiate_project_checklist_path", os.path.join(phase_output_dir, "Initiate_Project_Checklist.xlsx"))
        else:
            logging.warning("Không tìm thấy nội dung Initiate Project Checklist.")

    risk_tasks = Task(
        description=dedent(f"""
            Là {getattr(initiation_agent, 'role', 'Chuyên gia Khởi tạo')}, bạn hãy tiến hành đánh giá rủi ro sơ bộ cho dự án.
            Xác định các rủi ro tiềm ẩn, phân tích tác động và khả năng xảy ra của chúng.
            Tạo một tài liệu đánh giá rủi ro (Risk Assessment Document) chi tiết.
            Đồng thời, tạo một danh sách kiểm tra khởi tạo dự án (Initiate Project Checklist) cơ bản
            để đảm bảo tất cả các bước quan trọng đã được hoàn thành.
            Output của bạn phải chứa 2 phần riêng biệt, mỗi phần được bao bọc bởi các tag độc đáo:
            ```risk_assessment
            [Nội dung đánh giá rủi ro]
            ```
            ```initiate_project_checklist
            [Nội dung danh sách kiểm tra (định dạng bảng)]
            ```
            --- Context:
            Báo cáo Khả thi: {shared_memory.get('phase_0_initiation', 'feasibility_report_path') or 'N/A'}
            Khái niệm Vận hành (ConOps): {shared_memory.get('phase_0_initiation', 'concept_of_operations_path') or 'N/A'}
        """),
        expected_output="2 tài liệu tiếng Việt: 'Risk_Assessment_Document.md', 'Initiate_Project_Checklist.xlsx'.",
        agent=initiation_agent,
        context=[initiation_tasks, conops_tasks],
        callback=risk_tasks_callback
    )

    def estimate_tasks_callback(output):
        logging.info(f"--- Hoàn thành Task 'estimate_tasks' ---")
        prelim_schedule_match = re.search(r'```preliminary_schedule\s*([\s\S]*?)\s*```', str(output))
        budget_estimate_match = re.search(r'```budget_estimate\s*([\s\S]*?)\s*```', str(output))
        cost_benefit_analysis_match = re.search(r'```cost_benefit_analysis\s*([\s\S]*?)\s*```', str(output))
        if prelim_schedule_match:
            process_and_save_xlsx(prelim_schedule_match.group(1), os.path.join(phase_output_dir, "Preliminary_Schedule.xlsx"), "Lịch trình Sơ bộ")
            shared_memory.set("phase_0_initiation", "preliminary_schedule_path", os.path.join(phase_output_dir, "Preliminary_Schedule.xlsx"))
        else:
            logging.warning("Không tìm thấy nội dung Preliminary Schedule.")
        if budget_estimate_match:
            process_and_save_xlsx(budget_estimate_match.group(1), os.path.join(phase_output_dir, "Budget_Estimate.xlsx"), "Ước tính Ngân sách")
            shared_memory.set("phase_0_initiation", "budget_estimate_path", os.path.join(phase_output_dir, "Budget_Estimate.xlsx"))
        else:
            logging.warning("Không tìm thấy nội dung Budget Estimate.")
        if cost_benefit_analysis_match:
            process_and_save_docx(cost_benefit_analysis_match.group(1), os.path.join(phase_output_dir, "Cost_Benefit_Analysis.docx"), "Phân tích Chi phí-Lợi ích")
            shared_memory.set("phase_0_initiation", "cost_benefit_analysis_path", os.path.join(phase_output_dir, "Cost_Benefit_Analysis.docx"))
        else:
            logging.warning("Không tìm thấy nội dung Cost Benefit Analysis.")

    estimate_tasks = Task(
        description=dedent(f"""
            Là {getattr(initiation_agent, 'role', 'Chuyên gia Khởi tạo')}, bạn hãy tạo các ước tính ban đầu về lịch trình và ngân sách cho dự án.
            Phát triển một lịch trình sơ bộ (Preliminary Schedule) ở định dạng bảng (phù hợp cho Excel)
            với các mốc quan trọng và ước tính thời gian.
            Tạo một ước tính ngân sách (Budget Estimate) ở định dạng bảng (phù hợp cho Excel)
            bao gồm các hạng mục chi phí chính.
            Thực hiện phân tích chi phí-lợi ích (Cost-Benefit Analysis) để hỗ trợ quyết định.
            Output của bạn phải chứa 3 phần riêng biệt, mỗi phần được bao bọc bởi các tag độc đáo:
            ```preliminary_schedule
            [Nội dung lịch trình sơ bộ (định dạng bảng)]
            ```
            ```budget_estimate
            [Nội dung ước tính ngân sách (định dạng bảng)]
            ```
            ```cost_benefit_analysis
            [Nội dung phân tích chi phí-lợi ích]
            ```
            --- Context:
            Điều lệ Dự án: {shared_memory.get('phase_0_initiation', 'project_charter_path') or 'N/A'}
            Đánh giá Rủi ro: {shared_memory.get('phase_0_initiation', 'risk_assessment_document_path') or 'N/A'}
        """),
        expected_output="3 tài liệu tiếng Việt: 'Preliminary_Schedule.xlsx', 'Budget_Estimate.xlsx', 'Cost_Benefit_Analysis.docx'.",
        agent=initiation_agent,
        context=[initiation_tasks, risk_tasks],
        callback=estimate_tasks_callback
    )

    def quality_gate_initiation_task_callback(output):
        logging.info(f"--- Hoàn thành Project Initiation Validation Task ---")
        filepath = os.path.join(phase_output_dir, "Validation_Report_Phase_0.md")
        write_output(filepath, str(output))
        shared_memory.set("phase_0_initiation", "validation_report_path", filepath)
        logging.info(f"Đã lưu Validation_Report_Phase_0.md và cập nhật shared_memory.")

    quality_gate_initiation_task = Task(
        description=dedent(f"""
            Là Project Manager, bạn hãy thực hiện kiểm tra chất lượng (Quality Gate) cho Giai đoạn Khởi tạo (Phase 0).
            Đánh giá kỹ lưỡng tất cả các deliverables đã tạo ra trong giai đoạn này:
            - Project_Charter.docx
            - Business_Case.docx
            - Feasibility_Report.docx
            - Query_Keywords.txt
            - Phase_Specific_Requirements.docx
            - Best_Practices_Report.md
            - IEEE_UML_Guidelines.md
            - Architecture_References.md
            - Template_Library_Update.md
            - Stakeholder_List.xlsx
            - Stakeholder_Analysis.md
            - Project_Submission_Form.docx
            - Project_Team_Definition.docx
            - Project_Resource_Plan.xlsx
            - Concept_of_Operations.docx
            - Risk_Assessment_Document.md
            - Initiate_Project_Checklist.xlsx
            - Preliminary_Schedule.xlsx
            - Budget_Estimate.xlsx
            - Cost_Benefit_Analysis.docx
            Kiểm tra tính đầy đủ, chính xác, nhất quán và sự phù hợp với mục tiêu dự án ban đầu và yêu cầu người dùng.
            Tạo một báo cáo 'Validation_Report_Phase_0.md' tóm tắt kết quả đánh giá.
            Báo cáo này cần nêu rõ các điểm mạnh, các điểm cần cải thiện (nếu có), và xác nhận xem
            giai đoạn Khởi tạo đã hoàn thành đạt yêu cầu để chuyển sang giai đoạn tiếp theo hay chưa.
            --- Context:
            Yêu cầu Người dùng Ban đầu: {shared_memory.get('phase_0_initiation', 'system_request_summary') or 'N/A'}
        """),
        expected_output="Tài liệu tiếng Việt 'Validation_Report_Phase_0.md' tóm tắt kết quả đánh giá giai đoạn khởi tạo.",
        agent=project_manager_agent,
        context=[
            initiation_tasks,
            keyword_generation_tasks,
            phase_req_analysis_tasks,
            research_tasks,
            stakeholder_tasks,
            resourcing_tasks,
            conops_tasks,
            risk_tasks,
            estimate_tasks
        ],
        callback=quality_gate_initiation_task_callback
    )

    return [
        initiation_tasks,
        keyword_generation_tasks,
        phase_req_analysis_tasks,
        research_tasks,
        stakeholder_tasks,
        resourcing_tasks,
        conops_tasks,
        risk_tasks,
        estimate_tasks,
        quality_gate_initiation_task
    ]